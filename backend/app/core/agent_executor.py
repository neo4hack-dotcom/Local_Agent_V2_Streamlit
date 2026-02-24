from __future__ import annotations

import json
import re
from collections import deque
from datetime import datetime, timedelta, timezone
from email.utils import parsedate_to_datetime
from html import unescape
from pathlib import Path
from typing import Any
from urllib.parse import parse_qs, quote, unquote, urljoin, urlparse
from xml.etree import ElementTree as ET

import requests

from .db_connectors import connector_for
from .langgraph_agent import LangGraphAgentRunner
from .llm_client import LLMClient
from .models import AgentConfig, AgentType, DatabaseProfile, LLMConfig, agent_requires_database
from .web_navigation_agent import WebNavigationRunner

_JSON_FENCE = re.compile(r"```(?:json)?\\s*(.*?)```", flags=re.IGNORECASE | re.DOTALL)
_SQL_FENCE = re.compile(r"```(?:sql)?\\s*(.*?)```", flags=re.IGNORECASE | re.DOTALL)


class AgentExecutor:
    def __init__(self, llm_config: LLMConfig) -> None:
        self.llm_config = llm_config
        self.llm = LLMClient(llm_config)
        self.sql_runner = LangGraphAgentRunner()

    def execute(
        self,
        agent: AgentConfig,
        question: str,
        database: DatabaseProfile | None,
    ) -> dict[str, Any]:
        if agent_requires_database(agent.agent_type) and not database:
            raise ValueError(
                f"Agent '{agent.name}' requires a database, but none is configured/selected."
            )

        if agent.agent_type == "sql_analyst":
            return self._run_sql_analyst(agent, question, database)
        if agent.agent_type == "clickhouse_table_manager":
            return self._run_clickhouse_table_manager(agent, question, database)
        if agent.agent_type == "unstructured_to_structured":
            return self._run_unstructured_extractor(agent, question)
        if agent.agent_type == "email_cleaner":
            return self._run_email_cleaner(agent, question)
        if agent.agent_type == "file_assistant":
            return self._run_file_assistant(agent, question)
        if agent.agent_type == "text_file_manager":
            return self._run_text_file_manager(agent, question)
        if agent.agent_type == "excel_manager":
            return self._run_excel_manager(agent, question)
        if agent.agent_type == "word_manager":
            return self._run_word_manager(agent, question)
        if agent.agent_type == "elasticsearch_retriever":
            return self._run_elasticsearch_retriever(agent, question, database)
        if agent.agent_type == "internet_search":
            return self._run_internet_search(agent, question)
        if agent.agent_type == "rss_news":
            return self._run_rss_news(agent, question)
        if agent.agent_type == "web_scraper":
            return self._run_web_scraper(agent, question)
        if agent.agent_type == "web_navigator":
            return self._run_web_navigator(agent, question)
        if agent.agent_type == "wikipedia_retriever":
            return self._run_wikipedia_retriever(agent, question)
        if agent.agent_type == "rag_context":
            return self._run_rag_context(agent, question)

        raise ValueError(f"Unsupported agent type: {agent.agent_type}")

    def _run_sql_analyst(
        self,
        agent: AgentConfig,
        question: str,
        database: DatabaseProfile | None,
    ) -> dict[str, Any]:
        if not database:
            raise ValueError("Database is required for SQL analyst execution.")
        if database.engine not in {"clickhouse", "oracle"}:
            raise ValueError(
                "SQL analyst supports only ClickHouse/Oracle profiles. "
                f"Selected profile '{database.name}' uses engine '{database.engine}'."
            )

        cfg = agent.template_config if isinstance(agent.template_config, dict) else {}
        sql_mode = str(cfg.get("sql_use_case_mode", "llm_sql")).strip().lower()
        sql_template = str(cfg.get("sql_query_template", "")).strip()
        if sql_mode in {"parameterized_template", "templated_sql", "use_case"} and sql_template:
            return self._run_sql_parameterized_use_case(
                agent=agent,
                question=question,
                database=database,
            )

        output = self.sql_runner.run(
            question=question,
            agent=agent,
            database=database,
            llm=self.llm_config,
        )
        return {
            "sql": output.get("sql", ""),
            "rows": output.get("rows", []),
            "answer": output.get("answer", ""),
            "details": {
                "agent_type": agent.agent_type,
                "database_id": database.id,
                "database_name": database.name,
            },
        }

    def _run_sql_parameterized_use_case(
        self,
        agent: AgentConfig,
        question: str,
        database: DatabaseProfile,
    ) -> dict[str, Any]:
        cfg = agent.template_config if isinstance(agent.template_config, dict) else {}
        sql_template = str(cfg.get("sql_query_template", "")).strip()
        if not sql_template:
            raise ValueError(
                "SQL use-case mode requires template_config.sql_query_template."
            )

        parameter_specs = self._normalize_sql_use_case_parameters(cfg.get("sql_parameters"))
        if not parameter_specs:
            placeholders = self._extract_sql_template_placeholders(sql_template)
            parameter_specs = [
                {
                    "name": name,
                    "description": "",
                    "type": "string",
                    "required": True,
                    "format_hint": "",
                    "example": "",
                    "default_value": None,
                }
                for name in placeholders
            ]

        parameter_values = self._extract_sql_use_case_values(
            question=question,
            parameter_specs=parameter_specs,
        )
        rendered_sql, missing_required = self._render_sql_use_case_template(
            sql_template=sql_template,
            parameter_specs=parameter_specs,
            parameter_values=parameter_values,
            database_engine=database.engine,
        )
        if missing_required:
            raise ValueError(
                "Missing required SQL parameter(s): "
                + ", ".join(missing_required)
                + ". Update your request with these values."
            )

        connector = connector_for(database)
        rows = connector.run_query(rendered_sql, limit=agent.max_rows)
        rows_payload = json.dumps(rows, ensure_ascii=False)
        prompt = agent.answer_prompt_template.format(
            question=question,
            sql=rendered_sql,
            rows=rows_payload,
        )
        answer = self.llm.generate(prompt, system_prompt=agent.system_prompt)

        return {
            "sql": rendered_sql,
            "rows": rows,
            "answer": answer,
            "details": {
                "agent_type": agent.agent_type,
                "database_id": database.id,
                "database_name": database.name,
                "sql_mode": "parameterized_template",
                "sql_template": sql_template,
                "parameters_extracted": parameter_values,
            },
        }

    def _run_clickhouse_table_manager(
        self,
        agent: AgentConfig,
        question: str,
        database: DatabaseProfile | None,
    ) -> dict[str, Any]:
        if not database:
            raise ValueError("Database is required for ClickHouse table manager execution.")
        if database.engine != "clickhouse":
            raise ValueError(
                "ClickHouse table manager requires a ClickHouse profile. "
                f"Selected profile '{database.name}' uses engine '{database.engine}'."
            )

        cfg = agent.template_config if isinstance(agent.template_config, dict) else {}
        policy = {
            "protect_existing_tables": self._to_bool(
                cfg.get("protect_existing_tables"), default=True
            ),
            "allow_row_inserts": self._to_bool(cfg.get("allow_row_inserts"), default=True),
            "allow_row_updates": self._to_bool(cfg.get("allow_row_updates"), default=True),
            "allow_row_deletes": self._to_bool(cfg.get("allow_row_deletes"), default=False),
        }
        max_statements = self._to_int(
            cfg.get("max_statements"),
            default=8,
            minimum=1,
            maximum=40,
        )
        stop_on_error = self._to_bool(cfg.get("stop_on_error"), default=True)
        preview_limit = self._to_int(
            cfg.get("preview_select_rows"),
            default=min(agent.max_rows, 100),
            minimum=1,
            maximum=500,
        )

        connector = connector_for(database)
        schema = connector.schema_snapshot(
            agent.allowed_tables if isinstance(agent.allowed_tables, list) else None
        )
        planned_operations = self._plan_clickhouse_table_operations(
            question=question,
            schema=schema,
            max_statements=max_statements,
            policy=policy,
            system_prompt=agent.system_prompt,
        )
        operations = self._validate_clickhouse_operations(
            operations=planned_operations,
            policy=policy,
            max_statements=max_statements,
        )

        execution_rows: list[dict[str, Any]] = []
        executed_sql: list[str] = []
        for index, operation in enumerate(operations, start=1):
            sql = str(operation.get("sql", "")).strip().rstrip(";")
            purpose = str(operation.get("purpose", "")).strip()
            if not sql:
                continue

            try:
                result = connector.run_statement(sql, limit=preview_limit)
                statement_rows = result.get("rows")
                if not isinstance(statement_rows, list):
                    statement_rows = []
                row_preview = statement_rows[:preview_limit]
                execution_rows.append(
                    {
                        "index": index,
                        "status": "success",
                        "purpose": purpose or None,
                        "statement": sql,
                        "statement_type": result.get("statement_type", "command"),
                        "row_count": result.get("row_count"),
                        "message": result.get("message", ""),
                        "rows_preview": row_preview,
                    }
                )
                executed_sql.append(sql)
            except Exception as exc:  # noqa: BLE001
                error_message = str(exc)
                execution_rows.append(
                    {
                        "index": index,
                        "status": "failed",
                        "purpose": purpose or None,
                        "statement": sql,
                        "error": error_message,
                    }
                )
                if stop_on_error:
                    raise ValueError(
                        f"Statement #{index} failed: {error_message}"
                    ) from exc

        if not execution_rows:
            raise ValueError("No executable SQL statement was produced for this request.")

        sql_text = ";\n".join(executed_sql)
        rows_payload = json.dumps(execution_rows, ensure_ascii=False)
        prompt = agent.answer_prompt_template.format(
            question=question,
            sql=sql_text,
            rows=rows_payload,
        )
        answer = self.llm.generate(prompt, system_prompt=agent.system_prompt)

        return {
            "sql": sql_text,
            "rows": execution_rows,
            "answer": answer,
            "details": {
                "agent_type": agent.agent_type,
                "database_id": database.id,
                "database_name": database.name,
                "policy": policy,
                "max_statements": max_statements,
                "stop_on_error": stop_on_error,
                "planned_statements": len(planned_operations),
                "executed_statements": len(executed_sql),
            },
        }

    def _plan_clickhouse_table_operations(
        self,
        *,
        question: str,
        schema: str,
        max_statements: int,
        policy: dict[str, Any],
        system_prompt: str,
    ) -> list[dict[str, str]]:
        policy_lines = [
            f"- protect_existing_tables={bool(policy.get('protect_existing_tables', True))}",
            f"- allow_row_inserts={bool(policy.get('allow_row_inserts', True))}",
            f"- allow_row_updates={bool(policy.get('allow_row_updates', True))}",
            f"- allow_row_deletes={bool(policy.get('allow_row_deletes', False))}",
            f"- max_statements={max_statements}",
        ]

        prompt = (
            "Generate ClickHouse SQL operations for this request.\\n"
            "Return strict JSON only, with this schema:\\n"
            "{\\n"
            '  "operations": [\\n'
            '    {"sql": "SQL statement", "purpose": "short reason"}\\n'
            "  ]\\n"
            "}\\n"
            "Rules:\\n"
            "- Use ClickHouse SQL syntax only.\\n"
            "- Return between 1 and max_statements statements.\\n"
            "- Prefer CREATE TABLE IF NOT EXISTS for table creation.\\n"
            "- If safety policy forbids an operation, do not generate it.\\n"
            "- Keep each statement standalone (no markdown, no comments).\\n\\n"
            f"Safety policy:\\n{chr(10).join(policy_lines)}\\n\\n"
            f"Schema snapshot:\\n{schema}\\n\\n"
            f"User request:\\n{question}"
        )

        raw = self.llm.generate(prompt, system_prompt=system_prompt)
        operations = self._extract_sql_operations(raw, max_statements=max_statements)
        if not operations:
            raise ValueError(
                "The ClickHouse table manager could not produce executable SQL statements."
            )
        return operations

    def _extract_sql_operations(
        self,
        raw_text: str,
        *,
        max_statements: int,
    ) -> list[dict[str, str]]:
        parsed = self._parse_json_output(raw_text)
        operations: list[dict[str, str]] = []

        if isinstance(parsed, dict):
            payload_operations = parsed.get("operations")
            if isinstance(payload_operations, list):
                for item in payload_operations:
                    if isinstance(item, dict):
                        sql = str(item.get("sql", "")).strip()
                        purpose = str(item.get("purpose", "")).strip()
                    else:
                        sql = str(item).strip()
                        purpose = ""
                    if sql:
                        split_sql = self._split_sql_statements(sql) or [sql]
                        for statement in split_sql:
                            operations.append({"sql": statement, "purpose": purpose})
            elif isinstance(parsed.get("sql"), str):
                for sql in self._split_sql_statements(str(parsed.get("sql", ""))):
                    operations.append({"sql": sql, "purpose": ""})

        if not operations and isinstance(parsed, list):
            for item in parsed:
                if isinstance(item, dict):
                    sql = str(item.get("sql", "")).strip()
                    purpose = str(item.get("purpose", "")).strip()
                else:
                    sql = str(item).strip()
                    purpose = ""
                if sql:
                    split_sql = self._split_sql_statements(sql) or [sql]
                    for statement in split_sql:
                        operations.append({"sql": statement, "purpose": purpose})

        if not operations:
            cleaned = raw_text.strip()
            fence_match = _SQL_FENCE.search(cleaned)
            if fence_match:
                cleaned = fence_match.group(1).strip()
            for sql in self._split_sql_statements(cleaned):
                operations.append({"sql": sql, "purpose": ""})

        return operations[:max_statements]

    def _validate_clickhouse_operations(
        self,
        *,
        operations: list[dict[str, str]],
        policy: dict[str, Any],
        max_statements: int,
    ) -> list[dict[str, str]]:
        if len(operations) > max_statements:
            raise ValueError(
                f"Generated {len(operations)} SQL statements, exceeding max_statements={max_statements}."
            )

        validated: list[dict[str, str]] = []
        for operation in operations:
            sql = str(operation.get("sql", "")).strip().rstrip(";")
            if not sql:
                continue
            self._validate_clickhouse_statement_policy(sql=sql, policy=policy)
            validated.append(
                {
                    "sql": sql,
                    "purpose": str(operation.get("purpose", "")).strip(),
                }
            )

        if not validated:
            raise ValueError("No valid SQL statements remained after policy validation.")

        return validated

    def _validate_clickhouse_statement_policy(
        self,
        *,
        sql: str,
        policy: dict[str, Any],
    ) -> None:
        normalized = re.sub(r"\s+", " ", sql).strip().lower()
        if not normalized:
            raise ValueError("Empty SQL statement.")

        first_token = normalized.split(" ", 1)[0]
        protect_existing_tables = bool(policy.get("protect_existing_tables", True))
        allow_row_inserts = bool(policy.get("allow_row_inserts", True))
        allow_row_updates = bool(policy.get("allow_row_updates", True))
        allow_row_deletes = bool(policy.get("allow_row_deletes", False))

        always_blocked = {"grant", "revoke", "attach", "detach", "system", "kill"}
        if first_token in always_blocked:
            raise ValueError(f"Operation '{first_token}' is always blocked for safety.")

        if first_token in {"drop", "truncate", "rename"} and protect_existing_tables:
            raise ValueError(
                "Blocked by safety policy: dropping/truncating/renaming tables is disabled "
                "(protect_existing_tables=true)."
            )

        if first_token == "create":
            if not re.search(r"^\s*create\s+table\b", normalized):
                raise ValueError(
                    "Only CREATE TABLE statements are supported for this agent."
                )
            if protect_existing_tables and " or replace " in f" {normalized} ":
                raise ValueError(
                    "Blocked by safety policy: CREATE OR REPLACE is disabled "
                    "(protect_existing_tables=true)."
                )
            return

        if first_token == "insert":
            if not allow_row_inserts:
                raise ValueError(
                    "Blocked by safety policy: INSERT is disabled (allow_row_inserts=false)."
                )
            return

        if first_token == "alter":
            if re.search(r"\bupdate\b", normalized):
                if not allow_row_updates:
                    raise ValueError(
                        "Blocked by safety policy: row updates are disabled "
                        "(allow_row_updates=false)."
                    )
                return
            if re.search(r"\bdelete\b", normalized):
                if not allow_row_deletes:
                    raise ValueError(
                        "Blocked by safety policy: row deletes are disabled "
                        "(allow_row_deletes=false)."
                    )
                return
            if protect_existing_tables:
                raise ValueError(
                    "Blocked by safety policy: ALTER schema operations are disabled "
                    "(protect_existing_tables=true)."
                )
            return

        if first_token == "delete":
            if not allow_row_deletes:
                raise ValueError(
                    "Blocked by safety policy: DELETE is disabled (allow_row_deletes=false)."
                )
            return

        if first_token == "update":
            if not allow_row_updates:
                raise ValueError(
                    "Blocked by safety policy: UPDATE is disabled (allow_row_updates=false)."
                )
            return

        if first_token == "select":
            return

        if first_token == "with":
            if re.search(
                r"\b(insert|alter|drop|truncate|rename|delete|update|create)\b",
                normalized,
            ):
                raise ValueError(
                    "WITH statements containing write or schema operations are not supported."
                )
            return

        raise ValueError(
            f"Unsupported statement type '{first_token}' for ClickHouse table manager."
        )

    @staticmethod
    def _split_sql_statements(sql_script: str) -> list[str]:
        statements: list[str] = []
        buffer: list[str] = []
        in_single = False
        in_double = False
        in_backtick = False

        for char in sql_script:
            if char == "'" and not in_double and not in_backtick:
                in_single = not in_single
                buffer.append(char)
                continue
            if char == '"' and not in_single and not in_backtick:
                in_double = not in_double
                buffer.append(char)
                continue
            if char == "`" and not in_single and not in_double:
                in_backtick = not in_backtick
                buffer.append(char)
                continue

            if char == ";" and not in_single and not in_double and not in_backtick:
                candidate = "".join(buffer).strip()
                if candidate:
                    statements.append(candidate)
                buffer = []
                continue

            buffer.append(char)

        trailing = "".join(buffer).strip()
        if trailing:
            statements.append(trailing)

        return statements

    def _normalize_sql_use_case_parameters(self, raw_value: Any) -> list[dict[str, Any]]:
        if not isinstance(raw_value, list):
            return []

        normalized: list[dict[str, Any]] = []
        seen_names: set[str] = set()
        allowed_types = {"string", "integer", "number", "boolean", "date"}

        for item in raw_value:
            if not isinstance(item, dict):
                continue

            safe_name = self._normalize_sql_parameter_name(item.get("name", ""))
            if not safe_name or safe_name in seen_names:
                continue

            value_type = str(item.get("type", "string")).strip().lower()
            if value_type not in allowed_types:
                value_type = "string"

            required = self._to_bool(item.get("required"), default=True)
            default_value = item.get("default_value")
            if isinstance(default_value, str):
                default_value = default_value.strip() or None

            normalized.append(
                {
                    "name": safe_name,
                    "description": str(item.get("description", "")).strip(),
                    "type": value_type,
                    "required": required,
                    "format_hint": str(item.get("format_hint", "")).strip(),
                    "example": str(item.get("example", "")).strip(),
                    "default_value": default_value,
                }
            )
            seen_names.add(safe_name)

        return normalized

    @staticmethod
    def _normalize_sql_parameter_name(raw_name: Any) -> str:
        base = str(raw_name or "").strip().lower()
        if not base:
            return ""
        base = re.sub(r"[^a-z0-9_]+", "_", base)
        base = re.sub(r"_+", "_", base).strip("_")
        if not base:
            return ""
        if re.fullmatch(r"\d+", base):
            base = f"p_{base}"
        if base[0].isdigit():
            base = f"p_{base}"
        return base

    @staticmethod
    def _extract_sql_template_placeholders(sql_template: str) -> list[str]:
        names: list[str] = []
        seen: set[str] = set()

        for match in re.finditer(r"{{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*}}", sql_template):
            candidate = AgentExecutor._normalize_sql_parameter_name(match.group(1))
            if candidate and candidate not in seen:
                seen.add(candidate)
                names.append(candidate)

        for match in re.finditer(r"(?<!{){\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*}(?!})", sql_template):
            candidate = AgentExecutor._normalize_sql_parameter_name(match.group(1))
            if candidate and candidate not in seen:
                seen.add(candidate)
                names.append(candidate)

        return names

    def _extract_sql_use_case_values(
        self,
        *,
        question: str,
        parameter_specs: list[dict[str, Any]],
    ) -> dict[str, Any]:
        if not parameter_specs:
            return {}

        schema = [
            {
                "name": item["name"],
                "description": item.get("description", ""),
                "type": item.get("type", "string"),
                "required": bool(item.get("required", True)),
                "format_hint": item.get("format_hint", ""),
                "example": item.get("example", ""),
            }
            for item in parameter_specs
        ]

        prompt = (
            "Extract SQL use-case parameter values from this user request.\\n"
            "Return JSON only with this schema:\\n"
            "{\\n"
            '  "values": {"parameter_name": "value"},\\n'
            '  "missing_required": ["parameter_name"]\\n'
            "}\\n"
            "Rules:\\n"
            "- Use parameter names exactly as provided.\\n"
            "- Do not invent values.\\n"
            "- If a required parameter is missing, list it in missing_required.\\n\\n"
            f"Parameters schema: {json.dumps(schema, ensure_ascii=False)}\\n"
            f"User request:\\n{question}\\n"
        )

        extracted_payload: dict[str, Any] = {}
        try:
            raw = self.llm.generate(
                prompt,
                system_prompt=(
                    "You extract structured parameter values reliably. Return strict JSON only."
                ),
            )
            parsed = self._parse_json_output(raw)
            if isinstance(parsed, dict):
                if isinstance(parsed.get("values"), dict):
                    extracted_payload = {
                        self._normalize_sql_parameter_name(str(key)): value
                        for key, value in parsed["values"].items()
                        if self._normalize_sql_parameter_name(str(key))
                    }
                else:
                    extracted_payload = {
                        self._normalize_sql_parameter_name(str(key)): value
                        for key, value in parsed.items()
                        if key not in {"missing_required"}
                        and self._normalize_sql_parameter_name(str(key))
                    }
        except Exception:  # noqa: BLE001
            extracted_payload = {}

        values: dict[str, Any] = {}
        for spec in parameter_specs:
            name = str(spec.get("name", "")).strip()
            if not name:
                continue

            candidate = extracted_payload.get(name)
            if self._is_empty_parameter_value(candidate):
                candidate = self._fallback_extract_sql_parameter(
                    question=question,
                    parameter_name=name,
                )
            if self._is_empty_parameter_value(candidate):
                candidate = spec.get("default_value")
            if self._is_empty_parameter_value(candidate):
                continue

            values[name] = self._coerce_sql_parameter_value(
                raw_value=candidate,
                value_type=str(spec.get("type", "string")),
                parameter_name=name,
            )

        return values

    @staticmethod
    def _fallback_extract_sql_parameter(question: str, parameter_name: str) -> str | None:
        safe_name = re.escape(parameter_name).replace("_", r"[\s_-]*")
        patterns = [
            rf"\b{safe_name}\b\s*(?:=|:|is|est)\s*[\"']?([^\"'\n,;]+)",
            rf"\b{safe_name}\b\s*[\"']([^\"']+)[\"']",
        ]
        for pattern in patterns:
            match = re.search(pattern, question, flags=re.IGNORECASE)
            if not match:
                continue
            value = str(match.group(1)).strip()
            if value:
                return value
        return None

    def _render_sql_use_case_template(
        self,
        *,
        sql_template: str,
        parameter_specs: list[dict[str, Any]],
        parameter_values: dict[str, Any],
        database_engine: str,
    ) -> tuple[str, list[str]]:
        specs_by_name = {str(item.get("name")): item for item in parameter_specs}
        placeholder_names = self._extract_sql_template_placeholders(sql_template)
        rendered_sql = sql_template
        missing_required: list[str] = []

        for name in placeholder_names:
            spec = specs_by_name.get(
                name,
                {
                    "name": name,
                    "type": "string",
                    "required": True,
                    "default_value": None,
                },
            )
            value = parameter_values.get(name)
            if self._is_empty_parameter_value(value):
                default_value = spec.get("default_value")
                if not self._is_empty_parameter_value(default_value):
                    value = self._coerce_sql_parameter_value(
                        raw_value=default_value,
                        value_type=str(spec.get("type", "string")),
                        parameter_name=name,
                    )
                    parameter_values[name] = value

            if self._is_empty_parameter_value(value):
                if self._to_bool(spec.get("required"), default=True):
                    if name not in missing_required:
                        missing_required.append(name)
                    continue
                replacement = "NULL"
            else:
                replacement = self._sql_parameter_literal(
                    value=value,
                    value_type=str(spec.get("type", "string")),
                    database_engine=database_engine,
                )

            rendered_sql = re.sub(
                rf"{{{{\s*{re.escape(name)}\s*}}}}",
                replacement,
                rendered_sql,
                flags=re.IGNORECASE,
            )
            rendered_sql = re.sub(
                rf"(?<!{{){{\s*{re.escape(name)}\s*}}(?!}})",
                replacement,
                rendered_sql,
                flags=re.IGNORECASE,
            )

        unresolved = self._extract_sql_template_placeholders(rendered_sql)
        for name in unresolved:
            if name not in missing_required:
                missing_required.append(name)

        return rendered_sql, missing_required

    def _coerce_sql_parameter_value(
        self,
        *,
        raw_value: Any,
        value_type: str,
        parameter_name: str,
    ) -> Any:
        normalized_type = str(value_type or "string").strip().lower()
        if raw_value is None:
            return None

        if normalized_type == "string" or normalized_type == "date":
            value = str(raw_value).strip()
            return value or None

        if normalized_type == "integer":
            try:
                return int(str(raw_value).strip())
            except (TypeError, ValueError) as exc:
                raise ValueError(
                    f"Parameter '{parameter_name}' must be an integer."
                ) from exc

        if normalized_type == "number":
            try:
                return float(str(raw_value).strip())
            except (TypeError, ValueError) as exc:
                raise ValueError(
                    f"Parameter '{parameter_name}' must be a number."
                ) from exc

        if normalized_type == "boolean":
            if isinstance(raw_value, bool):
                return raw_value
            token = str(raw_value).strip().lower()
            if token in {"1", "true", "yes", "y", "on"}:
                return True
            if token in {"0", "false", "no", "n", "off"}:
                return False
            raise ValueError(
                f"Parameter '{parameter_name}' must be a boolean."
            )

        value = str(raw_value).strip()
        return value or None

    @staticmethod
    def _sql_parameter_literal(value: Any, value_type: str, database_engine: str) -> str:
        normalized_type = str(value_type or "string").strip().lower()
        if value is None:
            return "NULL"
        if normalized_type == "integer":
            return str(int(value))
        if normalized_type == "number":
            number = float(value)
            if number.is_integer():
                return str(int(number))
            return str(number)
        if normalized_type == "boolean":
            return "1" if bool(value) else "0"

        text = str(value).replace("'", "''")
        return f"'{text}'"

    @staticmethod
    def _is_empty_parameter_value(value: Any) -> bool:
        if value is None:
            return True
        if isinstance(value, str):
            return not value.strip()
        return False

    def _run_unstructured_extractor(self, agent: AgentConfig, question: str) -> dict[str, Any]:
        schema = agent.template_config.get(
            "output_schema",
            {
                "summary": "string",
                "entities": [{"type": "string", "value": "string"}],
                "priority": "low|medium|high",
            },
        )
        strict_json = bool(agent.template_config.get("strict_json", True))

        prompt = (
            "Transform the following unstructured text into structured JSON.\\n"
            f"Schema: {json.dumps(schema, ensure_ascii=False)}\\n"
            f"Input text:\\n{question}\\n"
        )
        if strict_json:
            prompt += "Return JSON only."

        raw = self.llm.generate(prompt, system_prompt=agent.system_prompt)
        structured = self._parse_json_output(raw)

        if structured is None:
            answer = (
                "The extractor could not produce valid JSON. Raw output returned below.\\n\\n"
                f"{raw}"
            )
            return {
                "sql": "",
                "rows": [],
                "answer": answer,
                "details": {
                    "agent_type": agent.agent_type,
                    "structured": None,
                    "raw_output": raw,
                },
            }

        rows = [structured] if isinstance(structured, dict) else []
        answer = json.dumps(structured, ensure_ascii=False, indent=2)
        return {
            "sql": "",
            "rows": rows,
            "answer": answer,
            "details": {
                "agent_type": agent.agent_type,
                "structured": structured,
            },
        }

    def _run_email_cleaner(self, agent: AgentConfig, question: str) -> dict[str, Any]:
        max_bullets = int(agent.template_config.get("max_bullets", 8))
        include_sections = agent.template_config.get(
            "include_sections",
            ["summary", "action_items", "deadlines", "risks"],
        )
        sections_text = ", ".join(str(item) for item in include_sections)

        prompt = (
            "Clean the following email by removing noise and keeping essential information.\\n"
            f"Sections required: {sections_text}.\\n"
            f"Maximum bullets per section: {max_bullets}.\\n"
            f"Email content:\\n{question}"
        )
        answer = self.llm.generate(prompt, system_prompt=agent.system_prompt)
        return {
            "sql": "",
            "rows": [],
            "answer": answer,
            "details": {
                "agent_type": agent.agent_type,
                "sections": include_sections,
                "max_bullets": max_bullets,
            },
        }

    def _run_file_assistant(self, agent: AgentConfig, question: str) -> dict[str, Any]:
        file_context = self._retrieve_file_context(agent, question, rag_mode=False)
        prompt = (
            f"Question: {question}\\n"
            f"Context snippets:\\n{file_context['context_text']}\\n"
            "Answer using only this context. Cite file paths used."
        )
        answer = self.llm.generate(prompt, system_prompt=agent.system_prompt)
        return {
            "sql": "",
            "rows": file_context["rows"],
            "answer": answer,
            "details": {
                "agent_type": agent.agent_type,
                "folder_path": file_context["folder_path"],
                "selected_files": file_context["selected_files"],
                "retrieval_mode": file_context["retrieval_mode"],
                "matched_items": file_context["matched_items"],
            },
        }

    def _run_text_file_manager(self, agent: AgentConfig, question: str) -> dict[str, Any]:
        cfg = agent.template_config
        root = self._resolve_managed_root(cfg)
        action = self._parse_text_file_action(question=question, cfg=cfg)
        operation = self._normalize_text_operation(action.get("operation"))
        if not operation:
            raise ValueError("Missing operation for text file manager.")

        default_encoding = str(cfg.get("default_encoding", "utf-8")).strip() or "utf-8"
        allow_overwrite = self._to_bool(cfg.get("allow_overwrite"), default=True)
        max_chars = self._to_int(
            action.get("max_chars"),
            default=self._to_int(cfg.get("max_chars_read"), default=12000, minimum=200, maximum=500000),
            minimum=200,
            maximum=500000,
        )

        if operation == "list":
            recursive = self._to_bool(action.get("recursive"), default=False)
            pattern = str(action.get("pattern", "")).strip().lower()
            target_raw = str(action.get("file_path", "")).strip()

            target_dir = root
            if target_raw:
                resolved = self._resolve_safe_path(root, target_raw)
                if resolved.is_file():
                    relative = self._to_relative_path(root, resolved)
                    rows = [
                        {
                            "path": relative,
                            "size_bytes": resolved.stat().st_size,
                        }
                    ]
                    return {
                        "sql": "",
                        "rows": rows,
                        "answer": f"Found 1 file: {relative}",
                        "details": {
                            "agent_type": agent.agent_type,
                            "operation": operation,
                            "folder_path": str(root),
                        },
                    }
                target_dir = resolved

            if not target_dir.exists() or not target_dir.is_dir():
                raise ValueError(f"Directory does not exist: {target_dir}")

            iterator = target_dir.rglob("*") if recursive else target_dir.glob("*")
            rows: list[dict[str, Any]] = []
            for path in iterator:
                if not path.is_file():
                    continue
                relative = self._to_relative_path(root, path)
                if pattern and pattern not in relative.lower():
                    continue
                rows.append({"path": relative, "size_bytes": path.stat().st_size})
                if len(rows) >= 500:
                    break

            base_display = "." if target_dir == root else self._to_relative_path(root, target_dir)
            return {
                "sql": "",
                "rows": rows,
                "answer": f"Listed {len(rows)} file(s) under '{base_display}'.",
                "details": {
                    "agent_type": agent.agent_type,
                    "operation": operation,
                    "folder_path": str(root),
                    "recursive": recursive,
                    "pattern": pattern or None,
                },
            }

        file_path_raw = str(
            action.get("file_path") or cfg.get("default_file_path") or ""
        ).strip()
        if not file_path_raw:
            raise ValueError(
                "No file path provided. Set template_config.default_file_path or provide a target path in your request."
            )
        target = self._resolve_safe_path(root, file_path_raw)

        if operation == "read":
            if not target.exists() or not target.is_file():
                raise ValueError(f"File does not exist: {target}")
            content = target.read_text(encoding=default_encoding, errors="ignore")
            preview = content[:max_chars]
            if len(content) > max_chars:
                preview += "\n...[truncated]"
            relative = self._to_relative_path(root, target)
            return {
                "sql": "",
                "rows": [
                    {
                        "path": relative,
                        "content_preview": preview,
                        "full_length": len(content),
                    }
                ],
                "answer": preview if preview else "(empty file)",
                "details": {
                    "agent_type": agent.agent_type,
                    "operation": operation,
                    "folder_path": str(root),
                    "file_path": relative,
                    "encoding": default_encoding,
                },
            }

        content = action.get("content", "")
        content_text = str(content) if content is not None else ""
        target.parent.mkdir(parents=True, exist_ok=True)
        relative = self._to_relative_path(root, target)

        if operation == "create":
            if target.exists() and not allow_overwrite:
                raise ValueError(f"File already exists and overwrite is disabled: {relative}")
            if content_text:
                target.write_text(content_text, encoding=default_encoding)
            else:
                target.touch(exist_ok=True)
            return {
                "sql": "",
                "rows": [{"path": relative, "bytes_written": len(content_text.encode(default_encoding))}],
                "answer": f"Created file '{relative}'.",
                "details": {
                    "agent_type": agent.agent_type,
                    "operation": operation,
                    "folder_path": str(root),
                    "file_path": relative,
                },
            }

        if operation == "write":
            if target.exists() and not allow_overwrite:
                raise ValueError(f"Overwrite is disabled for file: {relative}")
            target.write_text(content_text, encoding=default_encoding)
            return {
                "sql": "",
                "rows": [{"path": relative, "bytes_written": len(content_text.encode(default_encoding))}],
                "answer": f"Wrote {len(content_text)} characters to '{relative}'.",
                "details": {
                    "agent_type": agent.agent_type,
                    "operation": operation,
                    "folder_path": str(root),
                    "file_path": relative,
                },
            }

        if operation == "append":
            prepend_newline = self._to_bool(action.get("prepend_newline"), default=False)
            prefix = ""
            if prepend_newline and target.exists() and target.stat().st_size > 0:
                prefix = "\n"
            with target.open("a", encoding=default_encoding) as handle:
                handle.write(prefix + content_text)
            return {
                "sql": "",
                "rows": [{"path": relative, "bytes_appended": len((prefix + content_text).encode(default_encoding))}],
                "answer": f"Appended {len(content_text)} characters to '{relative}'.",
                "details": {
                    "agent_type": agent.agent_type,
                    "operation": operation,
                    "folder_path": str(root),
                    "file_path": relative,
                },
            }

        raise ValueError(
            f"Unsupported text operation '{operation}'. Allowed: read, create, write, append, list."
        )

    def _run_excel_manager(self, agent: AgentConfig, question: str) -> dict[str, Any]:
        try:
            from openpyxl import Workbook, load_workbook
        except ImportError as exc:  # pragma: no cover - runtime dependency
            raise ValueError(
                "Excel support requires the 'openpyxl' package. Run backend setup/update dependencies."
            ) from exc

        cfg = agent.template_config
        root = self._resolve_managed_root(cfg)
        action = self._parse_excel_action(question=question, cfg=cfg)
        operation = self._normalize_excel_operation(action.get("operation"))
        if not operation:
            raise ValueError("Missing operation for excel manager.")

        workbook_default = str(cfg.get("workbook_path", "workbook.xlsx")).strip() or "workbook.xlsx"
        default_sheet = str(cfg.get("default_sheet", "Sheet1")).strip() or "Sheet1"
        action, operation = self._repair_excel_action_if_needed(
            question=question,
            action=action,
            operation=operation,
            default_workbook=workbook_default,
            default_sheet=default_sheet,
        )

        file_path_raw = str(action.get("file_path", workbook_default)).strip() or workbook_default
        workbook_path = self._resolve_safe_path(root, file_path_raw, required_suffix=".xlsx")
        max_rows_read = self._to_int(
            action.get("max_rows"),
            default=self._to_int(cfg.get("max_rows_read"), default=200, minimum=1, maximum=10000),
            minimum=1,
            maximum=10000,
        )
        auto_create_workbook = self._to_bool(cfg.get("auto_create_workbook"), default=True)

        def load_or_create_workbook():
            if workbook_path.exists():
                return load_workbook(filename=workbook_path)
            if not auto_create_workbook:
                raise ValueError(f"Workbook does not exist: {workbook_path}")
            workbook_path.parent.mkdir(parents=True, exist_ok=True)
            wb = Workbook()
            wb.active.title = default_sheet
            wb.save(workbook_path)
            return wb

        if operation == "create_workbook":
            workbook_path.parent.mkdir(parents=True, exist_ok=True)
            wb = Workbook()
            requested_sheet = str(action.get("sheet_name", default_sheet)).strip() or default_sheet
            wb.active.title = requested_sheet

            additional_sheets = action.get("sheets", [])
            if isinstance(additional_sheets, list):
                existing = {sheet.title for sheet in wb.worksheets}
                for sheet_name in additional_sheets:
                    candidate = str(sheet_name).strip()
                    if not candidate or candidate in existing:
                        continue
                    wb.create_sheet(title=candidate)
                    existing.add(candidate)

            first_sheet = wb[requested_sheet]
            headers = action.get("headers", [])
            if isinstance(headers, list) and headers:
                first_sheet.append([self._json_safe_cell_value(item) for item in headers])

            rows_input = action.get("rows", [])
            normalized_rows = self._normalize_excel_rows(rows_input)
            for row in normalized_rows:
                first_sheet.append(row)

            wb.save(workbook_path)
            relative = self._to_relative_path(root, workbook_path)
            return {
                "sql": "",
                "rows": [
                    {
                        "file_path": relative,
                        "sheet_name": requested_sheet,
                        "created": True,
                        "rows_written": len(normalized_rows),
                    }
                ],
                "answer": f"Workbook '{relative}' created with sheet '{requested_sheet}'.",
                "details": {
                    "agent_type": agent.agent_type,
                    "operation": operation,
                    "folder_path": str(root),
                    "workbook_path": relative,
                },
            }

        wb = load_or_create_workbook()
        requested_sheet = str(action.get("sheet_name", default_sheet)).strip() or default_sheet

        if requested_sheet in wb.sheetnames:
            ws = wb[requested_sheet]
        elif operation in {"append_rows", "set_cells"}:
            ws = wb.create_sheet(title=requested_sheet)
        else:
            raise ValueError(
                f"Sheet '{requested_sheet}' does not exist in workbook '{self._to_relative_path(root, workbook_path)}'."
            )

        if operation == "list_sheets":
            rows = [
                {
                    "sheet_name": sheet.title,
                    "max_row": sheet.max_row,
                    "max_column": sheet.max_column,
                }
                for sheet in wb.worksheets
            ]
            return {
                "sql": "",
                "rows": rows,
                "answer": f"Workbook contains {len(rows)} sheet(s).",
                "details": {
                    "agent_type": agent.agent_type,
                    "operation": operation,
                    "folder_path": str(root),
                    "workbook_path": self._to_relative_path(root, workbook_path),
                },
            }

        if operation == "read_sheet":
            raw_rows: list[list[Any]] = []
            for values in ws.iter_rows(min_row=1, max_row=max_rows_read, values_only=True):
                row_values = [self._json_safe_cell_value(item) for item in values]
                raw_rows.append(row_values)

            use_header = self._to_bool(action.get("use_header"), default=True)
            if use_header and raw_rows:
                header = [
                    str(value).strip() if value is not None and str(value).strip() else f"column_{index + 1}"
                    for index, value in enumerate(raw_rows[0])
                ]
                mapped_rows: list[dict[str, Any]] = []
                for row in raw_rows[1:]:
                    mapped_rows.append(
                        {
                            header[index]: row[index] if index < len(row) else None
                            for index in range(len(header))
                        }
                    )
                rows: list[dict[str, Any]] = mapped_rows if mapped_rows else [
                    {"row_index": 1, "values": raw_rows[0]}
                ]
            else:
                rows = [
                    {"row_index": index + 1, "values": values}
                    for index, values in enumerate(raw_rows)
                ]

            return {
                "sql": "",
                "rows": rows,
                "answer": f"Read {len(rows)} row(s) from sheet '{ws.title}'.",
                "details": {
                    "agent_type": agent.agent_type,
                    "operation": operation,
                    "folder_path": str(root),
                    "workbook_path": self._to_relative_path(root, workbook_path),
                    "sheet_name": ws.title,
                    "max_rows": max_rows_read,
                },
            }

        if operation == "append_rows":
            normalized_rows = self._resolve_excel_rows_for_append(action=action, worksheet=ws)
            if not normalized_rows:
                raise ValueError(
                    "append_rows requires at least one non-empty row. "
                    "Example: rows=[['Alice', 1200, '2026-02']] or rows=[{'name':'Alice','amount':1200}]."
                )
            for row in normalized_rows:
                ws.append(row)
            wb.save(workbook_path)
            return {
                "sql": "",
                "rows": [
                    {
                        "sheet_name": ws.title,
                        "appended_rows": len(normalized_rows),
                        "new_max_row": ws.max_row,
                    }
                ],
                "answer": f"Appended {len(normalized_rows)} row(s) to sheet '{ws.title}'.",
                "details": {
                    "agent_type": agent.agent_type,
                    "operation": operation,
                    "folder_path": str(root),
                    "workbook_path": self._to_relative_path(root, workbook_path),
                    "sheet_name": ws.title,
                },
            }

        if operation == "set_cells":
            cell_updates = self._resolve_excel_cells_for_update(action=action)
            if not cell_updates:
                raise ValueError(
                    "set_cells requires at least one valid cell update. "
                    "Example: cells=[{'cell':'B2','value':42}]."
                )

            applied_updates: list[dict[str, Any]] = []
            for item in cell_updates:
                cell_ref = str(item.get("cell", "")).strip().upper()
                value = self._json_safe_cell_value(item.get("value"))
                ws[cell_ref] = value
                applied_updates.append({"cell": cell_ref, "value": value})

            wb.save(workbook_path)
            return {
                "sql": "",
                "rows": applied_updates,
                "answer": f"Updated {len(applied_updates)} cell(s) in sheet '{ws.title}'.",
                "details": {
                    "agent_type": agent.agent_type,
                    "operation": operation,
                    "folder_path": str(root),
                    "workbook_path": self._to_relative_path(root, workbook_path),
                    "sheet_name": ws.title,
                },
            }

        raise ValueError(
            "Unsupported excel operation. Allowed: create_workbook, list_sheets, read_sheet, append_rows, set_cells."
        )

    def _run_word_manager(self, agent: AgentConfig, question: str) -> dict[str, Any]:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - runtime dependency
            raise ValueError(
                "Word support requires the 'python-docx' package. Run backend setup/update dependencies."
            ) from exc

        cfg = agent.template_config
        root = self._resolve_managed_root(cfg)
        action = self._parse_word_action(question=question, cfg=cfg)
        operation = self._normalize_word_operation(action.get("operation"))
        if not operation:
            raise ValueError("Missing operation for word manager.")

        default_document = str(cfg.get("document_path", "document.docx")).strip() or "document.docx"
        file_path_raw = str(action.get("file_path", default_document)).strip() or default_document
        document_path = self._resolve_safe_path(root, file_path_raw, required_suffix=".docx")
        max_paragraphs_read = self._to_int(
            action.get("max_paragraphs"),
            default=self._to_int(
                cfg.get("max_paragraphs_read"), default=80, minimum=1, maximum=5000
            ),
            minimum=1,
            maximum=5000,
        )
        auto_create_document = self._to_bool(cfg.get("auto_create_document"), default=True)
        allow_overwrite = self._to_bool(cfg.get("allow_overwrite"), default=True)

        if operation == "list_documents":
            recursive = self._to_bool(action.get("recursive"), default=False)
            pattern = str(action.get("pattern", "")).strip().lower()
            iterator = root.rglob("*.docx") if recursive else root.glob("*.docx")
            rows: list[dict[str, Any]] = []
            for path in iterator:
                if not path.is_file():
                    continue
                relative = self._to_relative_path(root, path)
                if pattern and pattern not in relative.lower():
                    continue
                rows.append(
                    {
                        "file_path": relative,
                        "size_bytes": path.stat().st_size,
                    }
                )
                if len(rows) >= 500:
                    break

            return {
                "sql": "",
                "rows": rows,
                "answer": f"Listed {len(rows)} Word document(s).",
                "details": {
                    "agent_type": agent.agent_type,
                    "operation": operation,
                    "folder_path": str(root),
                    "recursive": recursive,
                    "pattern": pattern or None,
                },
            }

        def load_or_create_document():
            if document_path.exists():
                return Document(document_path)
            if not auto_create_document:
                raise ValueError(f"Document does not exist: {document_path}")
            document_path.parent.mkdir(parents=True, exist_ok=True)
            doc = Document()
            doc.save(document_path)
            return doc

        if operation == "create_document":
            relative = self._to_relative_path(root, document_path)
            if document_path.exists() and not allow_overwrite:
                raise ValueError(f"Overwrite is disabled for existing document: {relative}")

            document_path.parent.mkdir(parents=True, exist_ok=True)
            doc = Document()
            title = str(action.get("title", "")).strip()
            paragraphs = self._normalize_word_paragraphs(
                action.get("paragraphs"),
                fallback_content=action.get("content"),
            )
            title, paragraphs = self._resolve_word_content(
                question=question,
                operation=operation,
                title=title,
                paragraphs=paragraphs,
            )
            if title:
                doc.add_heading(title, level=1)
            for paragraph in paragraphs:
                doc.add_paragraph(paragraph)

            doc.save(document_path)
            return {
                "sql": "",
                "rows": [
                    {
                        "file_path": relative,
                        "created": True,
                        "paragraphs_written": len(paragraphs),
                    }
                ],
                "answer": (
                    f"Created Word document '{relative}' with {len(paragraphs)} paragraph(s)."
                ),
                "details": {
                    "agent_type": agent.agent_type,
                    "operation": operation,
                    "folder_path": str(root),
                    "file_path": relative,
                    "title": title or None,
                },
            }

        if operation == "read_document":
            if not document_path.exists() or not document_path.is_file():
                raise ValueError(f"Document does not exist: {document_path}")
            doc = Document(document_path)
            paragraphs = [
                paragraph.text.strip()
                for paragraph in doc.paragraphs
                if paragraph.text and paragraph.text.strip()
            ]
            selected = paragraphs[:max_paragraphs_read]
            rows = [
                {"paragraph_index": index + 1, "text": text}
                for index, text in enumerate(selected)
            ]
            preview = "\n\n".join(
                [f"[{row['paragraph_index']}] {row['text']}" for row in rows]
            ).strip()
            if len(paragraphs) > max_paragraphs_read:
                preview = f"{preview}\n\n...[truncated]"

            return {
                "sql": "",
                "rows": rows,
                "answer": preview or "(empty document)",
                "details": {
                    "agent_type": agent.agent_type,
                    "operation": operation,
                    "folder_path": str(root),
                    "file_path": self._to_relative_path(root, document_path),
                    "paragraph_count": len(paragraphs),
                    "max_paragraphs_read": max_paragraphs_read,
                },
            }

        if operation == "append_paragraphs":
            doc = load_or_create_document()
            paragraphs = self._normalize_word_paragraphs(
                action.get("paragraphs"),
                fallback_content=action.get("content"),
            )
            _, paragraphs = self._resolve_word_content(
                question=question,
                operation=operation,
                title="",
                paragraphs=paragraphs,
            )
            if not paragraphs:
                raise ValueError(
                    "append_paragraphs requires non-empty 'paragraphs' or 'content'."
                )
            for paragraph in paragraphs:
                doc.add_paragraph(paragraph)
            doc.save(document_path)
            return {
                "sql": "",
                "rows": [
                    {
                        "file_path": self._to_relative_path(root, document_path),
                        "appended_paragraphs": len(paragraphs),
                    }
                ],
                "answer": (
                    f"Appended {len(paragraphs)} paragraph(s) to "
                    f"'{self._to_relative_path(root, document_path)}'."
                ),
                "details": {
                    "agent_type": agent.agent_type,
                    "operation": operation,
                    "folder_path": str(root),
                    "file_path": self._to_relative_path(root, document_path),
                },
            }

        if operation == "replace_text":
            doc = load_or_create_document()
            find_text = str(action.get("find", "")).strip()
            replace_text = str(action.get("replace", ""))
            if not find_text:
                raise ValueError("replace_text requires a non-empty 'find' field.")

            replacement_count = 0
            for paragraph in doc.paragraphs:
                if not paragraph.text or find_text not in paragraph.text:
                    continue
                replacement_count += paragraph.text.count(find_text)
                paragraph.text = paragraph.text.replace(find_text, replace_text)

            doc.save(document_path)
            relative = self._to_relative_path(root, document_path)
            if replacement_count == 0:
                answer = (
                    f"No matching text found in '{relative}' for token '{find_text}'."
                )
            else:
                answer = (
                    f"Replaced {replacement_count} occurrence(s) of '{find_text}' in '{relative}'."
                )
            return {
                "sql": "",
                "rows": [
                    {
                        "file_path": relative,
                        "find": find_text,
                        "replace": replace_text,
                        "replacements": replacement_count,
                    }
                ],
                "answer": answer,
                "details": {
                    "agent_type": agent.agent_type,
                    "operation": operation,
                    "folder_path": str(root),
                    "file_path": relative,
                    "replacements": replacement_count,
                },
            }

        raise ValueError(
            "Unsupported word operation. Allowed: create_document, read_document, append_paragraphs, replace_text, list_documents."
        )

    def _run_elasticsearch_retriever(
        self,
        agent: AgentConfig,
        question: str,
        database: DatabaseProfile | None,
    ) -> dict[str, Any]:
        cfg = agent.template_config

        profile_base_url = ""
        profile_api_key = ""
        profile_username = ""
        profile_password = ""
        profile_verify_ssl = True
        if database and database.engine == "elasticsearch":
            if database.host:
                if database.host.startswith("http://") or database.host.startswith("https://"):
                    profile_base_url = database.host.rstrip("/")
                else:
                    scheme = "https" if database.secure else "http"
                    profile_base_url = f"{scheme}://{database.host}:{database.port or 9200}"
            profile_username = database.username or ""
            profile_password = database.password or ""
            if isinstance(database.options, dict):
                profile_api_key = str(database.options.get("api_key", "")).strip()
                profile_verify_ssl = bool(database.options.get("verify_ssl", True))

        base_url = str(cfg.get("base_url", "")).rstrip("/") or profile_base_url
        index = str(cfg.get("index", "")).strip()
        if not index and database and database.database:
            index = database.database
        top_k = int(cfg.get("top_k", 5))
        fields = cfg.get("fields", ["*"])
        verify_ssl = bool(cfg.get("verify_ssl", profile_verify_ssl))

        if not base_url:
            raise ValueError("Elasticsearch base_url is required in template_config.")
        if not index:
            raise ValueError("Elasticsearch index is required in template_config.")

        headers: dict[str, str] = {"Content-Type": "application/json"}
        api_key = str(cfg.get("api_key", "")).strip() or profile_api_key
        if api_key:
            headers["Authorization"] = f"ApiKey {api_key}"

        auth = None
        username = str(cfg.get("username", "")).strip() or profile_username
        password = str(cfg.get("password", "")).strip() or profile_password
        if username:
            auth = (username, password)

        query = {
            "size": top_k,
            "query": {
                "simple_query_string": {
                    "query": question,
                    "fields": fields if isinstance(fields, list) and fields else ["*"],
                }
            },
        }

        url = f"{base_url}/{index}/_search"
        response = requests.post(
            url,
            json=query,
            headers=headers,
            auth=auth,
            verify=verify_ssl,
            timeout=self.llm_config.timeout_seconds,
        )
        response.raise_for_status()
        body = response.json()

        hits = body.get("hits", {}).get("hits", [])
        rows: list[dict[str, Any]] = []
        for hit in hits[:top_k]:
            source = hit.get("_source") if isinstance(hit, dict) else {}
            if not isinstance(source, dict):
                source = {"value": source}
            rows.append(
                {
                    "_id": hit.get("_id"),
                    "_score": hit.get("_score"),
                    "_source": source,
                }
            )

        prompt = (
            f"Question: {question}\\n"
            f"Elasticsearch hits (JSON): {json.dumps(rows, ensure_ascii=False)}\\n"
            "Provide a concise answer grounded in these hits. Mention when evidence is weak."
        )
        answer = self.llm.generate(prompt, system_prompt=agent.system_prompt)
        return {
            "sql": "",
            "rows": rows,
            "answer": answer,
            "details": {
                "agent_type": agent.agent_type,
                "index": index,
                "hit_count": len(rows),
                "database_profile_id": database.id if database and database.engine == "elasticsearch" else None,
                "database_profile_name": database.name if database and database.engine == "elasticsearch" else None,
            },
        }

    def _run_rag_context(self, agent: AgentConfig, question: str) -> dict[str, Any]:
        file_context = self._retrieve_file_context(agent, question, rag_mode=True)
        prompt = (
            f"Question: {question}\\n"
            f"Retrieved context chunks:\\n{file_context['context_text']}\\n"
            "Answer with business context and cite source paths in your response."
        )
        answer = self.llm.generate(prompt, system_prompt=agent.system_prompt)
        return {
            "sql": "",
            "rows": file_context["rows"],
            "answer": answer,
            "details": {
                "agent_type": agent.agent_type,
                "folder_path": file_context["folder_path"],
                "selected_files": file_context["selected_files"],
                "retrieval_mode": file_context["retrieval_mode"],
                "matched_items": file_context["matched_items"],
            },
        }

    def _run_internet_search(self, agent: AgentConfig, question: str) -> dict[str, Any]:
        cfg = agent.template_config
        top_k = max(1, min(10, int(cfg.get("top_k", 5))))
        region = str(cfg.get("region", "wt-wt")).strip() or "wt-wt"
        safe_search = str(cfg.get("safe_search", "moderate")).strip().lower() or "moderate"
        fallback_to_instant = bool(cfg.get("fallback_to_instant", True))

        rows = self._duckduckgo_web_results(
            query=question,
            top_k=top_k,
            region=region,
            safe_search=safe_search,
        )

        if not rows and fallback_to_instant:
            rows = self._duckduckgo_instant_results(question, top_k=top_k)

        if not rows:
            raise ValueError(
                "Internet search returned no results. Try a different query or adjust template_config."
            )

        prompt = (
            f"Question: {question}\\n"
            f"Web search results (JSON): {json.dumps(rows, ensure_ascii=False)}\\n"
            "Provide a concise answer grounded in these sources. Include source links."
        )
        answer = self.llm.generate(prompt, system_prompt=agent.system_prompt)
        return {
            "sql": "",
            "rows": rows,
            "answer": answer,
            "details": {
                "agent_type": agent.agent_type,
                "provider": "duckduckgo",
                "region": region,
                "safe_search": safe_search,
                "result_count": len(rows),
            },
        }

    def _run_rss_news(self, agent: AgentConfig, question: str) -> dict[str, Any]:
        cfg = agent.template_config
        raw_feed_urls = cfg.get("feed_urls", [])
        feed_urls: list[str] = []
        if isinstance(raw_feed_urls, list):
            feed_urls = [str(item).strip() for item in raw_feed_urls if str(item).strip()]
        elif isinstance(raw_feed_urls, str) and raw_feed_urls.strip():
            feed_urls = [item.strip() for item in raw_feed_urls.split(",") if item.strip()]

        if not feed_urls:
            raise ValueError(
                "RSS feed URLs are required. Configure template_config.feed_urls with at least one URL."
            )

        interests = [
            str(item).strip().lower()
            for item in (cfg.get("interests", []) if isinstance(cfg.get("interests"), list) else [])
            if str(item).strip()
        ]
        exclude_keywords = [
            str(item).strip().lower()
            for item in (
                cfg.get("exclude_keywords", [])
                if isinstance(cfg.get("exclude_keywords"), list)
                else []
            )
            if str(item).strip()
        ]
        include_general_if_no_match = self._to_bool(
            cfg.get("include_general_if_no_match"),
            default=True,
        )
        top_k = self._to_int(cfg.get("top_k"), default=5, minimum=1, maximum=20)
        max_items_per_feed = self._to_int(
            cfg.get("max_items_per_feed"),
            default=25,
            minimum=1,
            maximum=200,
        )
        hours_lookback = self._to_int(
            cfg.get("hours_lookback"),
            default=24,
            minimum=1,
            maximum=24 * 14,
        )
        language_hint = str(cfg.get("language_hint", "fr")).strip() or "fr"

        cutoff = datetime.now(timezone.utc) - timedelta(hours=hours_lookback)
        collected: list[dict[str, Any]] = []
        feed_errors: list[str] = []

        for feed_url in feed_urls:
            normalized_feed_url = self._normalize_http_url(feed_url)
            if not normalized_feed_url:
                feed_errors.append(f"{feed_url}: invalid URL.")
                continue

            try:
                items = self._fetch_rss_feed_items(
                    feed_url=normalized_feed_url,
                    max_items=max_items_per_feed,
                )
            except ValueError as exc:
                feed_errors.append(f"{normalized_feed_url}: {exc}")
                continue

            for item in items:
                title = str(item.get("title", "")).strip()
                link = str(item.get("url", "")).strip()
                summary = str(item.get("summary", "")).strip()
                source = str(item.get("source", "")).strip() or urlparse(normalized_feed_url).netloc
                published_at = str(item.get("published_at", "")).strip()
                published_dt = self._parse_datetime_value(published_at)

                if published_dt and published_dt < cutoff:
                    continue

                searchable_text = f"{title} {summary}".lower()
                if not searchable_text.strip():
                    continue

                if exclude_keywords and any(keyword in searchable_text for keyword in exclude_keywords):
                    continue

                matched_interests = [
                    keyword for keyword in interests if keyword and keyword in searchable_text
                ]
                if interests and not matched_interests and not include_general_if_no_match:
                    continue

                recency_bonus = 0
                if published_dt:
                    age_hours = max(
                        0,
                        int((datetime.now(timezone.utc) - published_dt).total_seconds() // 3600),
                    )
                    recency_bonus = max(0, 24 - min(24, age_hours))

                score = len(matched_interests) * 20 + recency_bonus
                collected.append(
                    {
                        "title": title,
                        "url": link,
                        "summary": summary,
                        "source": source,
                        "published_at": published_at,
                        "matched_interests": matched_interests,
                        "score": score,
                    }
                )

        deduped: list[dict[str, Any]] = []
        seen_keys: set[str] = set()
        for row in collected:
            key = (str(row.get("url", "")).strip() or str(row.get("title", "")).strip()).lower()
            if not key or key in seen_keys:
                continue
            seen_keys.add(key)
            deduped.append(row)

        def sort_key(item: dict[str, Any]) -> tuple[int, int]:
            score_value = int(item.get("score", 0))
            published_dt = self._parse_datetime_value(str(item.get("published_at", "")))
            timestamp = int(published_dt.timestamp()) if published_dt else 0
            return (score_value, timestamp)

        deduped.sort(key=sort_key, reverse=True)
        selected_rows = deduped[:top_k]

        if not selected_rows:
            error_tail = f" Feed issues: {' | '.join(feed_errors[:3])}" if feed_errors else ""
            raise ValueError(
                "RSS News agent found no relevant articles. "
                "Adjust interests/exclusions or increase lookback window."
                + error_tail
            )

        target_briefing = question.strip() or "Prepare a short breakfast news briefing."
        prompt = (
            f"User briefing request: {target_briefing}\\n"
            f"Preferred language: {language_hint}\\n"
            f"User interests: {json.dumps(interests, ensure_ascii=False)}\\n"
            f"Excluded keywords: {json.dumps(exclude_keywords, ensure_ascii=False)}\\n"
            f"Selected news items (JSON): {json.dumps(selected_rows, ensure_ascii=False)}\\n\\n"
            "Write a concise breakfast briefing with this structure:\\n"
            "1) Headline of the day (1 line)\\n"
            "2) Top updates (3 to 5 bullets)\\n"
            "3) Why it matters today (2 bullets)\\n"
            "4) Sources (bullet list with links)\\n"
            "Keep it factual and compact."
        )
        answer = self.llm.generate(prompt, system_prompt=agent.system_prompt)

        return {
            "sql": "",
            "rows": selected_rows,
            "answer": answer,
            "details": {
                "agent_type": agent.agent_type,
                "feed_count": len(feed_urls),
                "feed_errors": feed_errors[:6],
                "lookback_hours": hours_lookback,
                "top_k": top_k,
                "interest_count": len(interests),
                "language_hint": language_hint,
            },
        }

    def _fetch_rss_feed_items(self, feed_url: str, max_items: int) -> list[dict[str, Any]]:
        headers = {
            **self._internet_request_headers(),
            "Accept": (
                "application/rss+xml,application/atom+xml,application/xml,text/xml;q=0.9,*/*;q=0.8"
            ),
        }
        try:
            response = requests.get(
                feed_url,
                headers=headers,
                timeout=self.llm_config.timeout_seconds,
            )
            response.raise_for_status()
        except requests.RequestException as exc:
            raise ValueError(f"feed request failed: {exc}") from exc

        payload = response.content.strip()
        if not payload:
            return []

        try:
            root = ET.fromstring(payload)
        except ET.ParseError:
            text_payload = response.text.lstrip("\ufeff").strip()
            if not text_payload:
                return []
            try:
                root = ET.fromstring(text_payload)
            except ET.ParseError as exc:
                raise ValueError(f"invalid XML feed format: {exc}") from exc

        root_name = self._xml_local_name(root.tag)
        if root_name == "feed":
            return self._parse_atom_feed(root, feed_url, max_items)
        if root_name in {"rss", "rdf"}:
            return self._parse_rss_feed(root, feed_url, max_items)

        if any(self._xml_local_name(child.tag) == "entry" for child in list(root)):
            return self._parse_atom_feed(root, feed_url, max_items)
        if any(self._xml_local_name(child.tag) in {"channel", "item"} for child in list(root)):
            return self._parse_rss_feed(root, feed_url, max_items)

        raise ValueError(f"unsupported XML feed root tag '{root.tag}'")

    def _parse_rss_feed(
        self,
        root: ET.Element,
        feed_url: str,
        max_items: int,
    ) -> list[dict[str, Any]]:
        channel = next(
            (child for child in list(root) if self._xml_local_name(child.tag) == "channel"),
            root,
        )
        source_name = self._xml_child_text(channel, ("title",)) or urlparse(feed_url).netloc

        rows: list[dict[str, Any]] = []
        for child in list(channel):
            if self._xml_local_name(child.tag) != "item":
                continue

            title = self._xml_child_text(child, ("title",))
            summary = self._xml_child_text(
                child,
                ("description", "summary", "content", "encoded"),
            )
            summary = summary[:4000]

            link_value = self._xml_child_raw_text(child, ("link", "guid"))
            normalized_link = self._normalize_http_url(link_value)
            if not normalized_link and link_value:
                normalized_link = self._normalize_http_url(urljoin(feed_url, link_value))

            published_raw = self._xml_child_raw_text(
                child,
                ("pubdate", "published", "updated", "date"),
            )
            published_dt = self._parse_datetime_value(published_raw)
            published_at = published_dt.isoformat() if published_dt else published_raw

            if not title and not summary:
                continue

            rows.append(
                {
                    "title": title or "(untitled article)",
                    "url": normalized_link,
                    "summary": summary,
                    "source": source_name,
                    "published_at": published_at,
                }
            )
            if len(rows) >= max_items:
                break

        return rows

    def _parse_atom_feed(
        self,
        root: ET.Element,
        feed_url: str,
        max_items: int,
    ) -> list[dict[str, Any]]:
        source_name = self._xml_child_text(root, ("title",)) or urlparse(feed_url).netloc

        rows: list[dict[str, Any]] = []
        for entry in list(root):
            if self._xml_local_name(entry.tag) != "entry":
                continue

            title = self._xml_child_text(entry, ("title",))
            summary = self._xml_child_text(entry, ("summary", "content", "subtitle"))
            summary = summary[:4000]
            normalized_link = self._extract_atom_link(entry, base_url=feed_url)

            published_raw = self._xml_child_raw_text(entry, ("published", "updated", "date"))
            published_dt = self._parse_datetime_value(published_raw)
            published_at = published_dt.isoformat() if published_dt else published_raw

            if not title and not summary:
                continue

            rows.append(
                {
                    "title": title or "(untitled article)",
                    "url": normalized_link,
                    "summary": summary,
                    "source": source_name,
                    "published_at": published_at,
                }
            )
            if len(rows) >= max_items:
                break

        return rows

    def _extract_atom_link(self, entry: ET.Element, base_url: str) -> str:
        fallback = ""
        for child in list(entry):
            if self._xml_local_name(child.tag) != "link":
                continue

            href = str(child.attrib.get("href", "")).strip()
            rel = str(child.attrib.get("rel", "")).strip().lower()
            candidate = href or "".join(child.itertext()).strip()
            if not candidate:
                continue

            normalized = self._normalize_http_url(candidate)
            if not normalized:
                normalized = self._normalize_http_url(urljoin(base_url, candidate))
            if not normalized:
                continue

            if rel in {"", "alternate"}:
                return normalized
            if not fallback:
                fallback = normalized

        return fallback

    @staticmethod
    def _xml_local_name(tag: str) -> str:
        value = str(tag)
        if "}" in value:
            value = value.split("}", 1)[1]
        if ":" in value:
            value = value.split(":", 1)[1]
        return value.strip().lower()

    @classmethod
    def _xml_child_raw_text(
        cls,
        element: ET.Element,
        names: tuple[str, ...] | list[str],
    ) -> str:
        expected = {str(name).strip().lower() for name in names if str(name).strip()}
        for child in list(element):
            if cls._xml_local_name(child.tag) not in expected:
                continue
            raw = "".join(child.itertext()).strip()
            if raw:
                return raw
        return ""

    @classmethod
    def _xml_child_text(
        cls,
        element: ET.Element,
        names: tuple[str, ...] | list[str],
    ) -> str:
        return cls._clean_html(cls._xml_child_raw_text(element, names))

    @staticmethod
    def _parse_datetime_value(value: str) -> datetime | None:
        raw_value = str(value or "").strip()
        if not raw_value:
            return None

        try:
            parsed = parsedate_to_datetime(raw_value)
        except (TypeError, ValueError, IndexError):
            parsed = None
        if parsed is not None:
            if parsed.tzinfo is None:
                parsed = parsed.replace(tzinfo=timezone.utc)
            return parsed.astimezone(timezone.utc)

        iso_value = raw_value.replace("Z", "+00:00")
        try:
            parsed_iso = datetime.fromisoformat(iso_value)
        except ValueError:
            parsed_iso = None
        if parsed_iso is not None:
            if parsed_iso.tzinfo is None:
                parsed_iso = parsed_iso.replace(tzinfo=timezone.utc)
            return parsed_iso.astimezone(timezone.utc)

        for date_format in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
            try:
                parsed_fallback = datetime.strptime(raw_value, date_format)
            except ValueError:
                continue
            return parsed_fallback.replace(tzinfo=timezone.utc)
        return None

    def _run_web_scraper(self, agent: AgentConfig, question: str) -> dict[str, Any]:
        cfg = agent.template_config
        include_urls_from_question = bool(cfg.get("include_urls_from_question", True))
        search_fallback = bool(cfg.get("search_fallback", True))
        follow_links = bool(cfg.get("follow_links", False))
        same_domain_only = bool(cfg.get("same_domain_only", True))
        max_pages = max(1, min(20, int(cfg.get("max_pages", 3))))
        max_links_per_page = max(1, min(30, int(cfg.get("max_links_per_page", 10))))
        max_chars_per_page = max(500, min(30000, int(cfg.get("max_chars_per_page", 6000))))
        timeout_seconds = max(
            3,
            min(
                120,
                int(cfg.get("timeout_seconds", self.llm_config.timeout_seconds)),
            ),
        )

        start_urls: list[str] = []
        raw_start_urls = cfg.get("start_urls", [])
        if isinstance(raw_start_urls, list):
            start_urls = [str(item).strip() for item in raw_start_urls if str(item).strip()]
        elif isinstance(raw_start_urls, str) and raw_start_urls.strip():
            start_urls = [item.strip() for item in raw_start_urls.split(",") if item.strip()]

        if self._is_placeholder_start_urls(start_urls):
            start_urls = []

        discovered_urls = self._extract_urls_from_text(question) if include_urls_from_question else []
        inferred_domain_urls = (
            self._extract_urls_from_domain_mentions(question) if include_urls_from_question else []
        )
        candidate_urls = self._dedupe_urls([*start_urls, *discovered_urls, *inferred_domain_urls])

        raw_allowed_domains = cfg.get("allowed_domains", [])
        allowed_domains: list[str] = []
        if isinstance(raw_allowed_domains, list):
            allowed_domains = [str(item).strip().lower() for item in raw_allowed_domains if str(item).strip()]
        elif isinstance(raw_allowed_domains, str) and raw_allowed_domains.strip():
            allowed_domains = [
                item.strip().lower()
                for item in raw_allowed_domains.split(",")
                if item.strip()
            ]

        if self._is_placeholder_allowed_domains(allowed_domains):
            allowed_domains = []

        referenced_domains = self._domains_from_urls([*discovered_urls, *inferred_domain_urls])
        for domain in referenced_domains:
            if domain not in allowed_domains:
                allowed_domains.append(domain)

        search_fallback_used = False
        if not candidate_urls and search_fallback:
            search_fallback_used = True
            search_rows = self._duckduckgo_web_results(
                query=question,
                top_k=max(2, min(8, max_pages)),
                region=str(cfg.get("region", "wt-wt")).strip() or "wt-wt",
                safe_search=str(cfg.get("safe_search", "moderate")).strip().lower() or "moderate",
            )
            candidate_urls = self._dedupe_urls(
                [str(item.get("url", "")).strip() for item in search_rows if isinstance(item, dict)]
            )
            for domain in self._domains_from_urls(candidate_urls):
                if domain not in allowed_domains:
                    allowed_domains.append(domain)

        if not candidate_urls:
            raise ValueError(
                "No target URL found. Provide URL/domain in the request, configure start_urls, "
                "or enable search fallback with a clearer query."
            )

        for domain in self._domains_from_urls(candidate_urls):
            if domain not in allowed_domains:
                allowed_domains.append(domain)

        rows = self._web_scrape_results(
            initial_urls=candidate_urls,
            allowed_domains=allowed_domains,
            max_pages=max_pages,
            max_links_per_page=max_links_per_page,
            max_chars_per_page=max_chars_per_page,
            follow_links=follow_links,
            same_domain_only=same_domain_only,
            timeout_seconds=timeout_seconds,
        )

        successful_rows = [
            row
            for row in rows
            if row.get("status") == "ok"
            and (str(row.get("snippet", "")).strip() or str(row.get("title", "")).strip())
        ]
        if not successful_rows:
            errors = [str(row.get("error", "")).strip() for row in rows if row.get("error")]
            error_text = " | ".join([error for error in errors[:3] if error])
            raise ValueError(
                "Web scraper could not extract readable content from target URLs."
                + (f" Errors: {error_text}" if error_text else "")
            )

        prompt = (
            f"Question: {question}\\n"
            f"Scraped web pages (JSON): {json.dumps(successful_rows, ensure_ascii=False)}\\n"
            "Answer using only these pages, and include source URLs."
        )
        answer = self.llm.generate(prompt, system_prompt=agent.system_prompt)
        return {
            "sql": "",
            "rows": successful_rows,
            "answer": answer,
            "details": {
                "agent_type": agent.agent_type,
                "scraped_pages": len(successful_rows),
                "attempted_pages": len(rows),
                "follow_links": follow_links,
                "allowed_domains": allowed_domains,
                "resolved_candidate_urls": candidate_urls,
                "search_fallback_used": search_fallback_used,
            },
        }

    def _run_web_navigator(self, agent: AgentConfig, question: str) -> dict[str, Any]:
        cfg = agent.template_config
        runner = WebNavigationRunner(
            llm=self.llm,
            system_prompt=agent.system_prompt,
            max_content_chars=max(500, min(30000, int(cfg.get("capture_html_chars", 7000)))),
        )
        output = runner.run(task=question, config=cfg)
        return {
            "sql": "",
            "rows": output.get("rows", []),
            "answer": output.get("answer", ""),
            "details": {
                "agent_type": agent.agent_type,
                **(output.get("details", {}) if isinstance(output.get("details"), dict) else {}),
            },
        }

    def _run_wikipedia_retriever(self, agent: AgentConfig, question: str) -> dict[str, Any]:
        cfg = agent.template_config
        language = str(cfg.get("language", "en")).strip() or "en"
        top_k = max(1, min(10, int(cfg.get("top_k", 5))))
        summary_sentences = max(1, min(5, int(cfg.get("summary_sentences", 2))))

        rows = self._wikipedia_results(
            query=question,
            language=language,
            top_k=top_k,
            summary_sentences=summary_sentences,
        )

        if not rows:
            raise ValueError(
                "Wikipedia search returned no results. Try a different query or language."
            )

        prompt = (
            f"Question: {question}\\n"
            f"Wikipedia results (JSON): {json.dumps(rows, ensure_ascii=False)}\\n"
            "Provide a factual answer grounded in these pages. Include citation links."
        )
        answer = self.llm.generate(prompt, system_prompt=agent.system_prompt)
        return {
            "sql": "",
            "rows": rows,
            "answer": answer,
            "details": {
                "agent_type": agent.agent_type,
                "language": language,
                "result_count": len(rows),
            },
        }

    def _retrieve_file_context(
        self,
        agent: AgentConfig,
        question: str,
        rag_mode: bool,
    ) -> dict[str, Any]:
        cfg = agent.template_config
        folder_path = str(cfg.get("folder_path", "")).strip()
        if not folder_path:
            raise ValueError("template_config.folder_path is required for this agent type.")

        root = Path(folder_path).expanduser()
        if not root.exists() or not root.is_dir():
            raise ValueError(f"Configured folder does not exist or is not a directory: {root}")

        extensions = cfg.get("file_extensions", [".txt", ".md", ".json", ".csv"])
        if not isinstance(extensions, list) or not extensions:
            extensions = [".txt", ".md", ".json", ".csv"]
        extension_set = {str(ext).lower() for ext in extensions}

        max_files = int(cfg.get("max_files", 40))
        max_file_size_kb = int(cfg.get("max_file_size_kb", 400))
        max_bytes = max_file_size_kb * 1024

        question_tokens = self._tokenize(question)
        scored: list[tuple[int, Path, str]] = []
        fallback_pool: list[tuple[int, Path, str]] = []
        processed_files = 0

        candidate_paths = [path for path in root.rglob("*") if path.is_file()]
        for path in candidate_paths:
            if extension_set and path.suffix.lower() not in extension_set:
                continue
            if path.stat().st_size > max_bytes:
                continue

            try:
                text = path.read_text(encoding="utf-8", errors="ignore")
            except OSError:
                continue

            if not text.strip():
                continue
            processed_files += 1
            if processed_files > max_files:
                break

            if rag_mode:
                chunks = self._chunk_text(
                    text=text,
                    chunk_size=int(cfg.get("chunk_size", 1200)),
                    overlap=int(cfg.get("chunk_overlap", 150)),
                )
                for chunk in chunks:
                    score = self._score_text(question_tokens, chunk)
                    if score > 0:
                        scored.append((score, path, chunk))
                    elif len(fallback_pool) < max_files * 8:
                        fallback_pool.append((score, path, chunk))
            else:
                snippet = text[:3000]
                score = self._score_text(question_tokens, snippet)
                if score > 0:
                    scored.append((score, path, snippet))
                elif len(fallback_pool) < max_files * 8:
                    fallback_pool.append((score, path, snippet))

        retrieval_mode = "keyword_match"
        if not scored:
            if fallback_pool:
                scored = fallback_pool
                retrieval_mode = "fallback_no_overlap"
            else:
                extensions_text = ", ".join(sorted(extension_set)) if extension_set else "(any)"
                raise ValueError(
                    "No readable text content found in configured folder. "
                    f"folder_path={root}, file_extensions={extensions_text}, "
                    f"max_file_size_kb={max_file_size_kb}."
                )

        scored.sort(key=lambda item: item[0], reverse=True)
        top_k = int(cfg.get("top_k_chunks" if rag_mode else "top_k", 6))
        selected = scored[:top_k]

        rows: list[dict[str, Any]] = []
        context_parts: list[str] = []
        selected_files: list[str] = []

        for rank, (_, path, snippet) in enumerate(selected, start=1):
            source = str(path)
            selected_files.append(source)
            rows.append(
                {
                    "rank": rank,
                    "source": source,
                    "snippet": snippet,
                }
            )
            context_parts.append(f"[#{rank}] source={source}\\n{snippet}")

        return {
            "folder_path": str(root),
            "rows": rows,
            "context_text": "\\n\\n".join(context_parts),
            "selected_files": sorted(set(selected_files)),
            "retrieval_mode": retrieval_mode,
            "matched_items": len(selected),
        }

    def _parse_text_file_action(self, question: str, cfg: dict[str, Any]) -> dict[str, Any]:
        direct = self._parse_json_output(question)
        if isinstance(direct, dict) and direct.get("operation"):
            return direct

        default_file_path = str(cfg.get("default_file_path", "")).strip()
        parser_prompt = (
            "Convert this user request into a JSON action for text file operations.\\n"
            "Allowed operations: read, create, write, append, list.\\n"
            "Output JSON schema:\\n"
            "{\\n"
            '  "operation": "read|create|write|append|list",\\n'
            '  "file_path": "relative/path.txt",\\n'
            '  "content": "text to write or append (optional)",\\n'
            '  "recursive": false,\\n'
            '  "pattern": "optional filter for list",\\n'
            '  "max_chars": 12000\\n'
            "}\\n"
            f"Default file path: {default_file_path or '(none)'}\\n"
            f"User request:\\n{question}\\n"
            "Return JSON only."
        )
        raw = self.llm.generate(
            parser_prompt,
            system_prompt="You convert natural language file requests into strict JSON actions.",
        )
        parsed = self._parse_json_output(raw)
        if isinstance(parsed, dict) and parsed.get("operation"):
            return parsed
        raise ValueError(
            "Could not parse text file operation. Provide a clearer instruction or JSON action."
        )

    def _parse_excel_action(self, question: str, cfg: dict[str, Any]) -> dict[str, Any]:
        direct = self._parse_json_output(question)
        if isinstance(direct, dict) and direct.get("operation"):
            return direct

        default_workbook = str(cfg.get("workbook_path", "workbook.xlsx")).strip() or "workbook.xlsx"
        default_sheet = str(cfg.get("default_sheet", "Sheet1")).strip() or "Sheet1"
        parser_prompt = (
            "Convert this user request into a JSON action for Excel operations.\\n"
            "Allowed operations: create_workbook, list_sheets, read_sheet, append_rows, set_cells.\\n"
            "Output JSON schema:\\n"
            "{\\n"
            '  "operation": "create_workbook|list_sheets|read_sheet|append_rows|set_cells",\\n'
            '  "file_path": "relative/path.xlsx",\\n'
            '  "sheet_name": "Sheet1",\\n'
            '  "headers": ["col1", "col2"],\\n'
            '  "rows": [["value1", "value2"]],\\n'
            '  "cells": [{"cell":"A1","value":"..."}],\\n'
            '  "max_rows": 200\\n'
            "}\\n"
            f"Default workbook path: {default_workbook}\\n"
            f"Default sheet: {default_sheet}\\n"
            f"User request:\\n{question}\\n"
            "Important rules:\\n"
            "- Never return placeholder values like value1/value2/col1/col2/...\\n"
            "- If operation is append_rows, rows must contain at least one real row.\\n"
            "- If operation is set_cells, cells must contain at least one valid A1 cell reference.\\n"
            "- Keep the same language as the user request.\\n"
            "Return JSON only."
        )
        raw = self.llm.generate(
            parser_prompt,
            system_prompt="You convert natural language spreadsheet requests into strict JSON actions.",
        )
        parsed = self._parse_json_output(raw)
        if isinstance(parsed, dict) and parsed.get("operation"):
            return parsed
        raise ValueError(
            "Could not parse excel operation. Provide a clearer instruction or JSON action."
        )

    def _repair_excel_action_if_needed(
        self,
        *,
        question: str,
        action: dict[str, Any],
        operation: str,
        default_workbook: str,
        default_sheet: str,
    ) -> tuple[dict[str, Any], str]:
        normalized_operation = self._normalize_excel_operation(operation)
        if normalized_operation not in {"append_rows", "set_cells"}:
            return action, normalized_operation

        has_rows = bool(self._resolve_excel_rows_for_append(action=action, worksheet=None))
        has_cells = bool(self._resolve_excel_cells_for_update(action=action))
        if (normalized_operation == "append_rows" and has_rows) or (
            normalized_operation == "set_cells" and has_cells
        ):
            return action, normalized_operation

        prompt = (
            "Repair the Excel action JSON so it can be executed.\\n"
            "Return JSON only.\\n"
            "Allowed operations: append_rows, set_cells.\\n"
            "If user intent is row insertion, return append_rows with non-empty rows.\\n"
            "If user intent is cell modification, return set_cells with non-empty cells.\\n"
            "Never return placeholders (value1/value2/col1/...).\\n"
            "For set_cells, each item must use A1 notation (example: B2).\\n\\n"
            f"Default workbook path: {default_workbook}\\n"
            f"Default sheet: {default_sheet}\\n"
            f"Current operation guess: {normalized_operation}\\n"
            f"Current parsed action JSON: {json.dumps(action, ensure_ascii=False)}\\n"
            f"User request:\\n{question}\\n"
        )

        try:
            raw = self.llm.generate(
                prompt,
                system_prompt=(
                    "You repair spreadsheet action payloads. Return strict JSON only."
                ),
            )
        except Exception:  # noqa: BLE001
            return action, normalized_operation

        repaired = self._parse_json_output(raw)
        if not isinstance(repaired, dict):
            return action, normalized_operation

        repaired_operation = self._normalize_excel_operation(
            repaired.get("operation", normalized_operation)
        )
        if repaired_operation not in {"append_rows", "set_cells"}:
            repaired_operation = normalized_operation

        merged: dict[str, Any] = {
            **action,
            **repaired,
            "operation": repaired_operation,
        }
        if not str(merged.get("file_path", "")).strip():
            merged["file_path"] = default_workbook
        if not str(merged.get("sheet_name", "")).strip():
            merged["sheet_name"] = default_sheet

        if repaired_operation == "append_rows":
            if not self._resolve_excel_rows_for_append(action=merged, worksheet=None):
                return action, normalized_operation
        else:
            if not self._resolve_excel_cells_for_update(action=merged):
                return action, normalized_operation

        return merged, repaired_operation

    def _resolve_excel_rows_for_append(
        self,
        *,
        action: dict[str, Any],
        worksheet: Any | None,
    ) -> list[list[Any]]:
        rows_input = action.get("rows", [])
        if isinstance(rows_input, dict):
            rows_input = [rows_input]
        if not isinstance(rows_input, list):
            return []

        headers = self._worksheet_headers(worksheet) if worksheet is not None else []
        normalized_rows: list[list[Any]] = []

        for row in rows_input:
            row_values: list[Any] = []
            if isinstance(row, dict):
                if headers:
                    normalized_map = {
                        self._normalize_excel_header_key(key): self._json_safe_cell_value(value)
                        for key, value in row.items()
                    }
                    row_values = [
                        normalized_map.get(self._normalize_excel_header_key(header))
                        for header in headers
                    ]
                    if not self._excel_row_has_meaningful_value(row_values):
                        row_values = [self._json_safe_cell_value(value) for value in row.values()]
                else:
                    row_values = [self._json_safe_cell_value(value) for value in row.values()]
            elif isinstance(row, list):
                row_values = [self._json_safe_cell_value(value) for value in row]
            elif row is not None:
                row_values = [self._json_safe_cell_value(row)]

            if self._excel_row_has_meaningful_value(row_values):
                normalized_rows.append(row_values)

        return normalized_rows

    def _resolve_excel_cells_for_update(
        self,
        *,
        action: dict[str, Any],
    ) -> list[dict[str, Any]]:
        cells_input = action.get("cells")
        if cells_input is None:
            cells_input = action.get("updates")
        if isinstance(cells_input, dict):
            cells_input = [cells_input]
        if not isinstance(cells_input, list):
            return []

        updates: list[dict[str, Any]] = []
        seen_cells: set[str] = set()
        for item in cells_input:
            if not isinstance(item, dict):
                continue

            cell_ref = str(item.get("cell", "")).strip().upper()
            if not cell_ref:
                column = str(item.get("column", "")).strip().upper()
                row = str(item.get("row", "")).strip()
                if re.fullmatch(r"[A-Z]{1,3}", column) and re.fullmatch(r"[1-9][0-9]*", row):
                    cell_ref = f"{column}{row}"

            if not re.fullmatch(r"[A-Z]{1,3}[1-9][0-9]*", cell_ref):
                continue
            if cell_ref in seen_cells:
                continue

            if "value" in item:
                raw_value = item.get("value")
            elif "new_value" in item:
                raw_value = item.get("new_value")
            else:
                continue

            seen_cells.add(cell_ref)
            updates.append(
                {
                    "cell": cell_ref,
                    "value": self._json_safe_cell_value(raw_value),
                }
            )

        return updates

    @staticmethod
    def _normalize_excel_header_key(value: Any) -> str:
        return re.sub(r"[^a-z0-9]+", "", str(value or "").strip().lower())

    @staticmethod
    def _worksheet_headers(worksheet: Any | None) -> list[str]:
        if worksheet is None:
            return []
        try:
            first_row = next(worksheet.iter_rows(min_row=1, max_row=1, values_only=True), ())
        except Exception:  # noqa: BLE001
            return []

        headers: list[str] = []
        for item in first_row:
            text = str(item).strip() if item is not None else ""
            headers.append(text)
        if not any(headers):
            return []
        return headers

    @classmethod
    def _excel_row_has_meaningful_value(cls, row_values: list[Any]) -> bool:
        for value in row_values:
            if value is None:
                continue
            if isinstance(value, str):
                text = value.strip()
                if not text:
                    continue
                if cls._is_placeholder_excel_value(text):
                    continue
                return True
            else:
                return True
        return False

    @staticmethod
    def _is_placeholder_excel_value(value: str) -> bool:
        text = str(value or "").strip().lower()
        if not text:
            return True
        if text in {"...", "tbd", "todo", "sample", "example", "value", "text"}:
            return True
        if re.fullmatch(r"value\s*\d+", text):
            return True
        if re.fullmatch(r"col(?:umn)?\s*\d+", text):
            return True
        if text in {"value1", "value2", "col1", "col2"}:
            return True
        return False

    def _parse_word_action(self, question: str, cfg: dict[str, Any]) -> dict[str, Any]:
        direct = self._parse_json_output(question)
        if isinstance(direct, dict) and direct.get("operation"):
            return direct

        default_document = str(cfg.get("document_path", "document.docx")).strip() or "document.docx"
        parser_prompt = (
            "Convert this user request into a JSON action for Word document operations.\\n"
            "Allowed operations: create_document, read_document, append_paragraphs, replace_text, list_documents.\\n"
            "Output JSON schema:\\n"
            "{\\n"
            '  "operation": "create_document|read_document|append_paragraphs|replace_text|list_documents",\\n'
            '  "file_path": "relative/path.docx",\\n'
            '  "title": "optional title for create_document",\\n'
            '  "paragraphs": ["paragraph text 1", "paragraph text 2"],\\n'
            '  "content": "alternative text payload",\\n'
            '  "find": "text to replace",\\n'
            '  "replace": "replacement text",\\n'
            '  "recursive": false,\\n'
            '  "pattern": "optional list filter",\\n'
            '  "max_paragraphs": 80\\n'
            "}\\n"
            f"Default document path: {default_document}\\n"
            f"User request:\\n{question}\\n"
            "Important rules:\\n"
            "- Never return placeholder text like 'paragraph text 1', 'optional title', '...'.\\n"
            "- For create_document and append_paragraphs, include concrete final prose in "
            "'paragraphs' (or in 'content').\\n"
            "- Keep the same language as the user request.\\n"
            "Return JSON only."
        )
        raw = self.llm.generate(
            parser_prompt,
            system_prompt="You convert natural language Word document requests into strict JSON actions.",
        )
        parsed = self._parse_json_output(raw)
        if isinstance(parsed, dict) and parsed.get("operation"):
            return parsed
        raise ValueError(
            "Could not parse word operation. Provide a clearer instruction or JSON action."
        )

    def _resolve_word_content(
        self,
        *,
        question: str,
        operation: str,
        title: str,
        paragraphs: list[str],
    ) -> tuple[str, list[str]]:
        clean_title = str(title or "").strip()
        clean_paragraphs = [str(item).strip() for item in paragraphs if str(item).strip()]
        non_placeholder = [
            item for item in clean_paragraphs if not self._is_placeholder_word_text(item)
        ]

        # Keep explicit meaningful user content as-is.
        if non_placeholder:
            if clean_title and self._is_placeholder_word_text(clean_title):
                clean_title = ""
            return clean_title, non_placeholder

        generated_title, generated_paragraphs = self._generate_word_content_from_request(
            question=question,
            operation=operation,
        )

        if generated_paragraphs:
            final_title = clean_title
            if not final_title or self._is_placeholder_word_text(final_title):
                final_title = generated_title
            return final_title, generated_paragraphs

        if clean_title and self._is_placeholder_word_text(clean_title):
            clean_title = ""

        # Last-resort fallback keeps behavior deterministic even when generation fails.
        fallback = question.strip()
        if fallback:
            return clean_title, [fallback]
        return clean_title, []

    def _generate_word_content_from_request(
        self,
        *,
        question: str,
        operation: str,
    ) -> tuple[str, list[str]]:
        prompt = (
            "Generate concrete Word document content from the user request.\\n"
            "Return JSON only with this schema:\\n"
            "{\\n"
            '  "title": "short title",\\n'
            '  "paragraphs": ["paragraph 1", "paragraph 2", "paragraph 3"]\\n'
            "}\\n\\n"
            f"Operation: {operation}\\n"
            f"User request:\\n{question}\\n\\n"
            "Rules:\\n"
            "- No placeholders, no templates, no TODO markers.\\n"
            "- Paragraphs must be final prose directly usable in a document.\\n"
            "- Keep the same language as the user request.\\n"
            "- Generate 2 to 6 concise paragraphs."
        )
        try:
            raw = self.llm.generate(
                prompt,
                system_prompt=(
                    "You write clean, production-ready business prose. Return strict JSON only."
                ),
            )
        except Exception:  # noqa: BLE001
            return "", []

        parsed = self._parse_json_output(raw)
        if not isinstance(parsed, dict):
            return "", []

        title = str(parsed.get("title", "")).strip()
        paragraphs = self._normalize_word_paragraphs(
            parsed.get("paragraphs"),
            fallback_content=parsed.get("content"),
        )
        filtered_paragraphs = [
            paragraph
            for paragraph in paragraphs
            if paragraph and not self._is_placeholder_word_text(paragraph)
        ]
        return title, filtered_paragraphs[:8]

    @staticmethod
    def _is_placeholder_word_text(value: str) -> bool:
        text = str(value or "").strip().lower()
        if not text:
            return True
        if text in {
            "...",
            "tbd",
            "todo",
            "optional title",
            "optional title for create_document",
            "alternative text payload",
            "content",
            "text",
            "string",
        }:
            return True
        if re.fullmatch(r"paragraph\s*text\s*\d+", text):
            return True
        if text.startswith("example ") or text.startswith("sample "):
            return True
        if "lorem ipsum" in text:
            return True
        return False

    @staticmethod
    def _normalize_text_operation(raw_operation: Any) -> str:
        operation = str(raw_operation or "").strip().lower()
        aliases = {
            "open": "read",
            "read": "read",
            "create": "create",
            "new": "create",
            "write": "write",
            "edit": "write",
            "modify": "write",
            "update": "write",
            "append": "append",
            "add": "append",
            "list": "list",
            "ls": "list",
        }
        return aliases.get(operation, operation)

    @staticmethod
    def _normalize_excel_operation(raw_operation: Any) -> str:
        operation = str(raw_operation or "").strip().lower()
        aliases = {
            "create": "create_workbook",
            "new": "create_workbook",
            "create_workbook": "create_workbook",
            "list": "list_sheets",
            "list_sheets": "list_sheets",
            "read": "read_sheet",
            "open": "read_sheet",
            "read_sheet": "read_sheet",
            "append": "append_rows",
            "append_rows": "append_rows",
            "insert_rows": "append_rows",
            "update_cells": "set_cells",
            "set_cells": "set_cells",
        }
        return aliases.get(operation, operation)

    @staticmethod
    def _normalize_word_operation(raw_operation: Any) -> str:
        operation = str(raw_operation or "").strip().lower()
        aliases = {
            "create": "create_document",
            "new": "create_document",
            "create_document": "create_document",
            "write": "create_document",
            "overwrite": "create_document",
            "read": "read_document",
            "open": "read_document",
            "read_document": "read_document",
            "append": "append_paragraphs",
            "edit": "append_paragraphs",
            "append_paragraphs": "append_paragraphs",
            "replace": "replace_text",
            "modify": "replace_text",
            "update": "replace_text",
            "replace_text": "replace_text",
            "list": "list_documents",
            "ls": "list_documents",
            "list_documents": "list_documents",
        }
        return aliases.get(operation, operation)

    @staticmethod
    def _to_bool(value: Any, default: bool) -> bool:
        if isinstance(value, bool):
            return value
        if isinstance(value, str):
            normalized = value.strip().lower()
            if normalized in {"1", "true", "yes", "y", "on"}:
                return True
            if normalized in {"0", "false", "no", "n", "off"}:
                return False
            return default
        if value is None:
            return default
        return bool(value)

    @staticmethod
    def _to_int(value: Any, default: int, minimum: int, maximum: int) -> int:
        try:
            parsed = int(value)
        except (TypeError, ValueError):
            parsed = default
        return max(minimum, min(maximum, parsed))

    @staticmethod
    def _json_safe_cell_value(value: Any) -> Any:
        if isinstance(value, (str, int, float, bool)) or value is None:
            return value
        return str(value)

    @staticmethod
    def _normalize_excel_rows(raw_rows: Any) -> list[list[Any]]:
        if not isinstance(raw_rows, list):
            return []
        normalized: list[list[Any]] = []
        for row in raw_rows:
            if isinstance(row, list):
                normalized.append([AgentExecutor._json_safe_cell_value(item) for item in row])
            elif isinstance(row, dict):
                values = [AgentExecutor._json_safe_cell_value(item) for item in row.values()]
                normalized.append(values)
            elif row is not None:
                normalized.append([AgentExecutor._json_safe_cell_value(row)])
        return normalized

    @staticmethod
    def _normalize_word_paragraphs(
        raw_paragraphs: Any,
        fallback_content: Any = None,
    ) -> list[str]:
        paragraphs: list[str] = []
        if isinstance(raw_paragraphs, list):
            for item in raw_paragraphs:
                text = str(item).strip()
                if text:
                    paragraphs.append(text)
        elif isinstance(raw_paragraphs, str):
            text = raw_paragraphs.strip()
            if text:
                paragraphs.append(text)

        if paragraphs:
            return paragraphs

        content = str(fallback_content or "").strip()
        if not content:
            return []

        lines = [line.strip() for line in content.splitlines() if line.strip()]
        return lines if lines else [content]

    @staticmethod
    def _to_relative_path(root: Path, target: Path) -> str:
        try:
            relative = target.relative_to(root)
            text = str(relative)
            return text if text else "."
        except ValueError:
            return str(target)

    def _resolve_managed_root(self, cfg: dict[str, Any]) -> Path:
        folder_path = str(cfg.get("folder_path", "")).strip()
        if not folder_path:
            raise ValueError("template_config.folder_path is required for this agent type.")

        root = Path(folder_path).expanduser()
        auto_create_folder = self._to_bool(cfg.get("auto_create_folder"), default=True)
        if not root.exists():
            if not auto_create_folder:
                raise ValueError(f"Configured folder does not exist: {root}")
            root.mkdir(parents=True, exist_ok=True)
        if not root.is_dir():
            raise ValueError(f"Configured folder is not a directory: {root}")
        return root.resolve()

    def _resolve_safe_path(
        self,
        root: Path,
        raw_path: str,
        required_suffix: str | None = None,
    ) -> Path:
        candidate = Path(raw_path).expanduser()
        if not candidate.is_absolute():
            candidate = root / candidate
        if required_suffix:
            normalized_suffix = required_suffix.lower()
            if not candidate.suffix:
                candidate = candidate.with_suffix(normalized_suffix)
            elif candidate.suffix.lower() != normalized_suffix:
                raise ValueError(f"Expected a '{normalized_suffix}' file path: {candidate}")

        resolved = candidate.resolve()
        try:
            resolved.relative_to(root)
        except ValueError as exc:
            raise ValueError(
                f"Path '{resolved}' is outside the configured folder '{root}'."
            ) from exc
        return resolved

    def _duckduckgo_web_results(
        self,
        query: str,
        top_k: int,
        region: str,
        safe_search: str,
    ) -> list[dict[str, Any]]:
        safe_map = {"off": "-2", "moderate": "-1", "strict": "1"}
        params: dict[str, str] = {"q": query, "kl": region}
        if safe_search in safe_map:
            params["kp"] = safe_map[safe_search]

        try:
            response = requests.get(
                "https://html.duckduckgo.com/html/",
                params=params,
                headers=self._internet_request_headers(),
                timeout=self.llm_config.timeout_seconds,
            )
            response.raise_for_status()
        except requests.RequestException:
            return []

        html_text = response.text
        link_pattern = re.compile(
            r'<a[^>]*class="result__a"[^>]*href="([^"]+)"[^>]*>(.*?)</a>',
            flags=re.IGNORECASE | re.DOTALL,
        )
        snippet_pattern = re.compile(
            r'class="result__snippet"[^>]*>(.*?)</(?:a|div)>',
            flags=re.IGNORECASE | re.DOTALL,
        )

        snippets = [self._clean_html(snippet) for snippet in snippet_pattern.findall(html_text)]
        rows: list[dict[str, Any]] = []
        seen_urls: set[str] = set()

        for rank, match in enumerate(link_pattern.finditer(html_text), start=1):
            raw_href = match.group(1)
            title = self._clean_html(match.group(2))
            url = self._resolve_duckduckgo_href(raw_href)
            if not url or url in seen_urls:
                continue

            seen_urls.add(url)
            rows.append(
                {
                    "rank": len(rows) + 1,
                    "title": title,
                    "url": url,
                    "snippet": snippets[rank - 1] if rank - 1 < len(snippets) else "",
                }
            )
            if len(rows) >= top_k:
                break

        return rows

    def _duckduckgo_instant_results(self, query: str, top_k: int) -> list[dict[str, Any]]:
        try:
            response = requests.get(
                "https://api.duckduckgo.com/",
                params={
                    "q": query,
                    "format": "json",
                    "no_html": 1,
                    "no_redirect": 1,
                },
                headers=self._internet_request_headers(),
                timeout=self.llm_config.timeout_seconds,
            )
            response.raise_for_status()
            body = response.json()
        except requests.RequestException as exc:
            raise ValueError(f"Internet fallback search failed: {exc}") from exc

        rows: list[dict[str, Any]] = []
        abstract_text = str(body.get("AbstractText", "")).strip()
        abstract_url = str(body.get("AbstractURL", "")).strip()
        heading = str(body.get("Heading", "")).strip() or "DuckDuckGo instant result"
        if abstract_text and abstract_url:
            rows.append(
                {
                    "rank": 1,
                    "title": heading,
                    "url": abstract_url,
                    "snippet": abstract_text,
                }
            )

        def collect_related(items: list[Any]) -> None:
            for item in items:
                if len(rows) >= top_k:
                    return
                if isinstance(item, dict):
                    if "Topics" in item and isinstance(item["Topics"], list):
                        collect_related(item["Topics"])
                        continue
                    text = str(item.get("Text", "")).strip()
                    first_url = str(item.get("FirstURL", "")).strip()
                    if text and first_url:
                        rows.append(
                            {
                                "rank": len(rows) + 1,
                                "title": text.split(" - ", 1)[0][:140],
                                "url": first_url,
                                "snippet": text,
                            }
                        )

        related = body.get("RelatedTopics")
        if isinstance(related, list):
            collect_related(related)

        return rows[:top_k]

    def _wikipedia_results(
        self,
        query: str,
        language: str,
        top_k: int,
        summary_sentences: int,
    ) -> list[dict[str, Any]]:
        endpoint = f"https://{language}.wikipedia.org/w/api.php"
        headers = self._internet_request_headers()

        try:
            search_response = requests.get(
                endpoint,
                params={
                    "action": "query",
                    "list": "search",
                    "srsearch": query,
                    "srlimit": top_k,
                    "format": "json",
                    "utf8": 1,
                },
                headers=headers,
                timeout=self.llm_config.timeout_seconds,
            )
            search_response.raise_for_status()
            search_body = search_response.json()
        except requests.HTTPError as exc:
            status_code = exc.response.status_code if exc.response is not None else None
            if status_code in {403, 429}:
                fallback_rows = self._wikipedia_opensearch_results(
                    query=query,
                    language=language,
                    top_k=top_k,
                    summary_sentences=summary_sentences,
                )
                if fallback_rows:
                    return fallback_rows
            raise ValueError(f"Wikipedia search request failed: {exc}") from exc
        except requests.RequestException as exc:
            raise ValueError(f"Wikipedia search request failed: {exc}") from exc

        search_items = search_body.get("query", {}).get("search", [])
        if not isinstance(search_items, list) or not search_items:
            return []

        titles = [
            str(item.get("title", "")).strip()
            for item in search_items
            if isinstance(item, dict) and str(item.get("title", "")).strip()
        ][:top_k]
        if not titles:
            return []

        page_by_title: dict[str, dict[str, Any]] = {}
        try:
            details_response = requests.get(
                endpoint,
                params={
                    "action": "query",
                    "prop": "extracts",
                    "exintro": 1,
                    "explaintext": 1,
                    "redirects": 1,
                    "format": "json",
                    "titles": "|".join(titles),
                },
                headers=headers,
                timeout=self.llm_config.timeout_seconds,
            )
            details_response.raise_for_status()
            details_body = details_response.json()
            pages = details_body.get("query", {}).get("pages", {})
            if isinstance(pages, dict):
                for page in pages.values():
                    if isinstance(page, dict):
                        title = str(page.get("title", "")).strip()
                        if title:
                            page_by_title[title] = page
        except requests.RequestException:
            page_by_title = {}

        rows: list[dict[str, Any]] = []
        for rank, item in enumerate(search_items[:top_k], start=1):
            if not isinstance(item, dict):
                continue
            title = str(item.get("title", "")).strip()
            if not title:
                continue
            page = page_by_title.get(title, {})
            page_id = page.get("pageid")
            extract = str(page.get("extract", "")).strip()
            snippet = self._clean_html(str(item.get("snippet", "")))
            summary = self._first_sentences(extract, summary_sentences)
            if not summary:
                summary = snippet

            if isinstance(page_id, int):
                url = f"https://{language}.wikipedia.org/?curid={page_id}"
            else:
                url = f"https://{language}.wikipedia.org/wiki/{quote(title.replace(' ', '_'))}"

            rows.append(
                {
                    "rank": rank,
                    "title": title,
                    "url": url,
                    "snippet": snippet,
                    "summary": summary,
                }
            )
        return rows

    def _wikipedia_opensearch_results(
        self,
        query: str,
        language: str,
        top_k: int,
        summary_sentences: int,
    ) -> list[dict[str, Any]]:
        endpoint = f"https://{language}.wikipedia.org/w/api.php"
        try:
            response = requests.get(
                endpoint,
                params={
                    "action": "opensearch",
                    "search": query,
                    "limit": top_k,
                    "namespace": 0,
                    "format": "json",
                },
                headers=self._internet_request_headers(),
                timeout=self.llm_config.timeout_seconds,
            )
            response.raise_for_status()
            payload = response.json()
        except requests.RequestException:
            return []

        if not isinstance(payload, list) or len(payload) < 4:
            return []

        raw_titles = payload[1] if isinstance(payload[1], list) else []
        raw_descriptions = payload[2] if isinstance(payload[2], list) else []
        raw_links = payload[3] if isinstance(payload[3], list) else []

        rows: list[dict[str, Any]] = []
        for index, raw_title in enumerate(raw_titles[:top_k]):
            title = str(raw_title).strip()
            if not title:
                continue
            description = (
                str(raw_descriptions[index]).strip()
                if index < len(raw_descriptions)
                else ""
            )
            url = str(raw_links[index]).strip() if index < len(raw_links) else ""
            if not url:
                url = f"https://{language}.wikipedia.org/wiki/{quote(title.replace(' ', '_'))}"
            rows.append(
                {
                    "rank": len(rows) + 1,
                    "title": title,
                    "url": url,
                    "snippet": description,
                    "summary": self._first_sentences(description, summary_sentences)
                    or description,
                }
            )
        return rows

    def _web_scrape_results(
        self,
        initial_urls: list[str],
        allowed_domains: list[str],
        max_pages: int,
        max_links_per_page: int,
        max_chars_per_page: int,
        follow_links: bool,
        same_domain_only: bool,
        timeout_seconds: int,
    ) -> list[dict[str, Any]]:
        queue: deque[str] = deque()
        seen_queue: set[str] = set()
        visited: set[str] = set()
        rows: list[dict[str, Any]] = []

        for raw_url in initial_urls:
            normalized = self._normalize_http_url(raw_url)
            if not normalized:
                continue
            if not self._is_url_allowed(normalized, allowed_domains):
                continue
            if normalized in seen_queue:
                continue
            queue.append(normalized)
            seen_queue.add(normalized)

        if not queue:
            return rows

        headers = self._browser_request_headers()
        with requests.Session() as session:
            while queue and len(rows) < max_pages:
                current_url = queue.popleft()
                if current_url in visited:
                    continue
                visited.add(current_url)

                try:
                    response = session.get(
                        current_url,
                        headers=headers,
                        timeout=timeout_seconds,
                    )
                    response.raise_for_status()
                except requests.RequestException as exc:
                    rows.append(
                        {
                            "rank": len(rows) + 1,
                            "url": current_url,
                            "title": "",
                            "snippet": "",
                            "status": "error",
                            "error": str(exc),
                        }
                    )
                    continue

                html_text = response.text
                title = self._extract_html_title(html_text)
                snippet = self._html_to_text(html_text)[:max_chars_per_page].strip()
                content_type = str(response.headers.get("content-type", "")).lower()

                rows.append(
                    {
                        "rank": len(rows) + 1,
                        "url": current_url,
                        "title": title,
                        "snippet": snippet,
                        "status": "ok",
                        "content_type": content_type,
                        "status_code": response.status_code,
                    }
                )

                if not follow_links:
                    continue

                current_domain = urlparse(current_url).netloc.lower()
                links = self._extract_html_links(
                    html=html_text,
                    base_url=current_url,
                    max_links=max_links_per_page,
                )
                for link in links:
                    if same_domain_only and urlparse(link).netloc.lower() != current_domain:
                        continue
                    if not self._is_url_allowed(link, allowed_domains):
                        continue
                    if link in seen_queue or link in visited:
                        continue
                    queue.append(link)
                    seen_queue.add(link)

        return rows

    @staticmethod
    def _extract_urls_from_text(text: str) -> list[str]:
        pattern = re.compile(r"https?://[^\s<>\"]+", flags=re.IGNORECASE)
        results: list[str] = []
        seen: set[str] = set()
        for match in pattern.findall(text):
            normalized = AgentExecutor._normalize_http_url(match)
            if not normalized or normalized in seen:
                continue
            seen.add(normalized)
            results.append(normalized)
        return results

    @staticmethod
    def _extract_urls_from_domain_mentions(text: str) -> list[str]:
        pattern = re.compile(
            r"\b((?:[a-z0-9](?:[a-z0-9-]{0,61}[a-z0-9])?\.)+[a-z]{2,}(?:/[^\s<>\"]*)?)",
            flags=re.IGNORECASE,
        )
        blocked_suffixes = {".json", ".txt", ".csv", ".md", ".log"}
        results: list[str] = []
        seen: set[str] = set()
        for raw_match in pattern.findall(text):
            candidate = str(raw_match).strip().strip("\"'()[]{}<>,;")
            if not candidate:
                continue
            lowered = candidate.lower()
            if "@" in lowered:
                continue
            if any(lowered.endswith(suffix) for suffix in blocked_suffixes):
                continue
            normalized = AgentExecutor._normalize_http_url(f"https://{candidate}")
            if not normalized or normalized in seen:
                continue
            seen.add(normalized)
            results.append(normalized)
        return results

    @staticmethod
    def _domains_from_urls(urls: list[str]) -> list[str]:
        domains: list[str] = []
        seen: set[str] = set()
        for raw_url in urls:
            host = urlparse(raw_url).netloc.lower().split(":", 1)[0].strip()
            if not host:
                continue
            if host in seen:
                continue
            seen.add(host)
            domains.append(host)
        return domains

    @staticmethod
    def _dedupe_urls(urls: list[str]) -> list[str]:
        deduped: list[str] = []
        seen: set[str] = set()
        for raw_url in urls:
            normalized = AgentExecutor._normalize_http_url(str(raw_url))
            if not normalized or normalized in seen:
                continue
            seen.add(normalized)
            deduped.append(normalized)
        return deduped

    @staticmethod
    def _is_placeholder_start_urls(urls: list[str]) -> bool:
        normalized = {
            AgentExecutor._normalize_http_url(str(url)) for url in urls if str(url).strip()
        }
        normalized.discard("")
        placeholders = {
            "https://example.com",
            "https://www.example.com",
            "http://example.com",
            "http://www.example.com",
        }
        return bool(normalized) and normalized.issubset(placeholders)

    @staticmethod
    def _is_placeholder_allowed_domains(domains: list[str]) -> bool:
        normalized = {str(domain).strip().lower().lstrip(".") for domain in domains if str(domain).strip()}
        placeholders = {"example.com", "www.example.com"}
        return bool(normalized) and normalized.issubset(placeholders)

    @staticmethod
    def _normalize_http_url(value: str) -> str:
        candidate = value.strip().strip("\"'()[]{}<>,.;")
        if not candidate:
            return ""
        parsed = urlparse(candidate)
        if not parsed.scheme or not parsed.netloc:
            return ""
        if parsed.scheme.lower() not in {"http", "https"}:
            return ""
        return parsed._replace(fragment="").geturl()

    @staticmethod
    def _is_url_allowed(url: str, allowed_domains: list[str]) -> bool:
        if not allowed_domains:
            return True
        host = urlparse(url).netloc.lower().split(":", 1)[0]
        for domain in allowed_domains:
            normalized = domain.strip().lower().lstrip(".")
            if not normalized:
                continue
            if host == normalized or host.endswith(f".{normalized}"):
                return True
        return False

    @staticmethod
    def _extract_html_title(html: str) -> str:
        match = re.search(r"<title[^>]*>(.*?)</title>", html, flags=re.IGNORECASE | re.DOTALL)
        if not match:
            return ""
        return AgentExecutor._clean_html(match.group(1))

    @staticmethod
    def _html_to_text(html: str) -> str:
        cleaned = re.sub(
            r"<(script|style|noscript)[^>]*>.*?</\1>",
            " ",
            html,
            flags=re.IGNORECASE | re.DOTALL,
        )
        cleaned = re.sub(r"<br\s*/?>", "\n", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"</(p|div|li|h[1-6])>", "\n", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"<[^>]+>", " ", cleaned)
        cleaned = unescape(cleaned)
        cleaned = re.sub(r"\s+", " ", cleaned)
        return cleaned.strip()

    @staticmethod
    def _extract_html_links(html: str, base_url: str, max_links: int) -> list[str]:
        pattern = re.compile(r'href=["\']([^"\']+)["\']', flags=re.IGNORECASE)
        links: list[str] = []
        seen: set[str] = set()
        for raw_link in pattern.findall(html):
            href = unescape(raw_link).strip()
            if not href:
                continue
            lowered = href.lower()
            if lowered.startswith("#") or lowered.startswith("mailto:") or lowered.startswith(
                "javascript:"
            ):
                continue
            absolute = urljoin(base_url, href)
            normalized = AgentExecutor._normalize_http_url(absolute)
            if not normalized or normalized in seen:
                continue
            seen.add(normalized)
            links.append(normalized)
            if len(links) >= max_links:
                break
        return links

    @staticmethod
    def _internet_request_headers() -> dict[str, str]:
        return {
            "User-Agent": (
                "Local-Agent-Studio/1.0 "
                "(+https://github.com/neo4hack-dotcom/Local_Agent)"
            ),
            "Accept": "application/json,text/html;q=0.9,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.8,fr-FR;q=0.7,fr;q=0.6",
            "DNT": "1",
        }

    @staticmethod
    def _browser_request_headers() -> dict[str, str]:
        return {
            "User-Agent": (
                "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/124.0.0.0 Safari/537.36"
            ),
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.8,fr-FR;q=0.7,fr;q=0.6",
            "DNT": "1",
        }

    @staticmethod
    def _resolve_duckduckgo_href(raw_href: str) -> str:
        href = unescape(raw_href).strip()
        if not href:
            return ""
        if href.startswith("/l/?"):
            query = parse_qs(urlparse(href).query)
            uddg = query.get("uddg", [""])[0]
            if uddg:
                return unquote(uddg)
        if href.startswith("//"):
            return f"https:{href}"
        if href.startswith("/"):
            return f"https://duckduckgo.com{href}"
        return href

    @staticmethod
    def _clean_html(value: str) -> str:
        text = re.sub(r"<[^>]+>", " ", unescape(value))
        text = re.sub(r"\s+", " ", text)
        return text.strip()

    @staticmethod
    def _first_sentences(text: str, sentence_count: int) -> str:
        cleaned = re.sub(r"\s+", " ", text).strip()
        if not cleaned:
            return ""
        parts = re.split(r"(?<=[.!?])\s+", cleaned)
        return " ".join(parts[:sentence_count]).strip()

    @staticmethod
    def _tokenize(text: str) -> set[str]:
        return {token for token in re.findall(r"\b\w{3,}\b", text.lower(), flags=re.UNICODE)}

    def _score_text(self, question_tokens: set[str], text: str) -> int:
        if not question_tokens:
            return 0
        text_tokens = self._tokenize(text)
        return len(question_tokens & text_tokens)

    @staticmethod
    def _chunk_text(text: str, chunk_size: int, overlap: int) -> list[str]:
        if chunk_size <= 0:
            chunk_size = 1000
        if overlap < 0:
            overlap = 0
        if overlap >= chunk_size:
            overlap = chunk_size // 4

        chunks: list[str] = []
        start = 0
        step = max(1, chunk_size - overlap)
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            start += step
        return chunks

    @staticmethod
    def _parse_json_output(raw_text: str) -> Any | None:
        cleaned = raw_text.strip()
        match = _JSON_FENCE.search(cleaned)
        if match:
            cleaned = match.group(1).strip()

        try:
            return json.loads(cleaned)
        except json.JSONDecodeError:
            start = cleaned.find("{")
            end = cleaned.rfind("}")
            if start != -1 and end != -1 and end > start:
                try:
                    return json.loads(cleaned[start : end + 1])
                except json.JSONDecodeError:
                    return None
            return None
